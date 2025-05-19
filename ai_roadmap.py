import os
import streamlit as st
import json
import pandas as pd
from openai import OpenAI
import io
import tempfile
from datetime import datetime

# Initialize the OpenAI client
# the newest OpenAI model is "gpt-4o" which was released May 13, 2024
# do not change this unless explicitly requested by the user
client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

def process_resume(file_content):
    """
    Extract information from a resume using OpenAI's API.
    
    Args:
        file_content (str): The content of the resume file
        
    Returns:
        dict: Extracted information about the person
    """
    try:
        prompt = f"""
        Analyze the following resume content and extract key information:
        
        {file_content}
        
        Please extract and structure the following information in JSON format:
        1. Skills (technical and soft skills)
        2. Experience (work history, roles, and accomplishments)
        3. Education (degrees, certifications, and relevant training)
        4. Projects (personal or professional projects of note)
        5. Strengths (apparent strengths based on the resume)
        6. Areas for growth (potential areas for skill development)
        
        Format your response as a valid JSON object with these categories as keys.
        """
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"}
        )
        
        # Parse the JSON response
        resume_data = json.loads(response.choices[0].message.content)
        return resume_data
    except Exception as e:
        st.error(f"Error processing resume: {str(e)}")
        return {}

def process_questionnaire(responses):
    """
    Process questionnaire responses using OpenAI's API.
    
    Args:
        responses (dict): The user's responses to the questionnaire
        
    Returns:
        dict: Analysis of the user's preferences and characteristics
    """
    try:
        # Convert the responses dictionary to a string
        responses_str = json.dumps(responses, indent=2)
        
        prompt = f"""
        Analyze the following career questionnaire responses:
        
        {responses_str}
        
        Please extract and structure the following insights in JSON format:
        1. Personality traits
        2. Work style preferences
        3. Career motivations
        4. Strengths
        5. Areas for growth
        6. Preferred work environment
        7. Management style (leader or individual contributor preference)
        8. Risk tolerance
        9. Time horizon for career goals
        10. Recommended career focus areas based on these responses
        
        Format your response as a valid JSON object with these categories as keys.
        """
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"}
        )
        
        # Parse the JSON response
        analysis = json.loads(response.choices[0].message.content)
        return analysis
    except Exception as e:
        st.error(f"Error processing questionnaire: {str(e)}")
        return {}

def generate_ai_roadmap(pathway, user_data, metrics_data):
    """
    Generate a personalized career roadmap using OpenAI's API.
    
    Args:
        pathway (dict): The selected career pathway
        user_data (dict): Information about the user (from resume and questionnaire)
        metrics_data (dict): Information about metrics
        
    Returns:
        dict: A personalized roadmap
    """
    try:
        # Convert complex data structures to strings for the prompt
        pathway_str = json.dumps(pathway, indent=2)
        user_data_str = json.dumps(user_data, indent=2)
        
        # Create a context-aware prompt for the AI
        prompt = f"""
        Create a personalized career roadmap for someone pursuing the following career pathway:
        
        PATHWAY INFORMATION:
        {pathway_str}
        
        USER INFORMATION:
        {user_data_str}
        
        Please create a detailed, actionable, and personalized career roadmap that:
        
        1. Takes into account the user's current skills, experience, and preferences
        2. Charts a realistic path from their current state to success in the chosen pathway
        3. Divides the journey into 4 phases:
           - Research & Preparation (0-6 months)
           - Initial Development (6-18 months)
           - Growth & Advancement (1.5-3 years)
           - Mastery & Scale (3-5+ years)
        4. For each phase, provide:
           - 3-5 specific action items tailored to both the pathway and the user's background
           - Resources to utilize (specific books, courses, communities, tools)
           - Goals/milestones to achieve
           - Potential challenges and how to overcome them
        5. Include personalized advice for leveraging their strengths and addressing development areas
        
        Format your response as a JSON object with the following structure:
        {{
          "pathway_name": string,
          "personalized_summary": string,
          "phases": [
            {{
              "name": string,
              "description": string,
              "timeline": string,
              "action_items": [
                {{
                  "title": string,
                  "description": string,
                  "rationale": string,
                  "resources": [string]
                }}
              ],
              "goals": [string],
              "challenges": [
                {{
                  "challenge": string,
                  "solution": string
                }}
              ]
            }}
          ],
          "personal_strengths_leverage": [string],
          "development_areas_advice": [string],
          "long_term_vision": string
        }}
        """
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
            max_tokens=2500
        )
        
        # Parse the JSON response
        roadmap = json.loads(response.choices[0].message.content)
        return roadmap
    except Exception as e:
        st.error(f"Error generating AI roadmap: {str(e)}")
        # Create a basic error roadmap
        return {
            "pathway_name": pathway.get('name', 'Unknown Pathway'),
            "personalized_summary": "Unable to generate a personalized roadmap at this time. Please try again later.",
            "phases": [],
            "personal_strengths_leverage": [],
            "development_areas_advice": [],
            "long_term_vision": ""
        }

def display_ai_roadmap(roadmap):
    """
    Display an AI-generated personalized roadmap in the Streamlit app.
    
    Args:
        roadmap (dict): The AI-generated roadmap to display
    """
    st.title(f"Your Personalized Roadmap for: {roadmap['pathway_name']}")
    
    # Show the personalized summary
    st.write("## Personalized Overview")
    st.write(roadmap["personalized_summary"])
    
    # Show strengths and development areas
    col1, col2 = st.columns(2)
    
    with col1:
        st.write("### Leveraging Your Strengths")
        for strength in roadmap.get("personal_strengths_leverage", []):
            st.markdown(f"- {strength}")
    
    with col2:
        st.write("### Development Areas")
        for area in roadmap.get("development_areas_advice", []):
            st.markdown(f"- {area}")
    
    # Display the long-term vision
    if "long_term_vision" in roadmap and roadmap["long_term_vision"]:
        st.write("### Your Long-Term Vision")
        st.write(roadmap["long_term_vision"])
    
    # Display each phase of the roadmap
    st.write("## Your Roadmap Phases")
    
    for i, phase in enumerate(roadmap.get("phases", [])):
        with st.expander(f"Phase {i+1}: {phase['name']} - {phase.get('timeline', '')}", expanded=(i==0)):
            st.write(f"**{phase.get('description', '')}**")
            
            # Display goals
            if "goals" in phase and phase["goals"]:
                st.write("### Goals & Milestones")
                for goal in phase["goals"]:
                    st.markdown(f"- {goal}")
            
            # Display action items
            if "action_items" in phase and phase["action_items"]:
                st.write("### Action Items")
                for item in phase["action_items"]:
                    st.markdown(f"#### {item.get('title', '')}")
                    st.write(item.get('description', ''))
                    
                    if "rationale" in item and item["rationale"]:
                        st.write(f"*Why this matters:* {item['rationale']}")
                    
                    if "resources" in item and item["resources"]:
                        st.write("**Resources:**")
                        for resource in item["resources"]:
                            st.markdown(f"- {resource}")
            
            # Display challenges and solutions
            if "challenges" in phase and phase["challenges"]:
                st.write("### Potential Challenges & Solutions")
                for challenge in phase["challenges"]:
                    st.markdown(f"**Challenge:** {challenge.get('challenge', '')}")
                    st.markdown(f"**Solution:** {challenge.get('solution', '')}")
    
    # Add download option
    roadmap_json = json.dumps(roadmap, indent=2)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    st.download_button(
        label="Download Your Personalized Roadmap (JSON)",
        data=roadmap_json,
        file_name=f"personalized_roadmap_{timestamp}.json",
        mime="application/json"
    )

def create_questionnaire():
    """
    Create a career questionnaire to gather information about the user.
    
    Returns:
        dict: The user's responses to the questionnaire
    """
    st.write("## Career Development Questionnaire")
    st.write("Answer the following questions to help us personalize your career roadmap.")
    
    # Initialize responses dictionary
    responses = {}
    
    # Background section
    st.write("### Background")
    responses["years_experience"] = st.slider(
        "How many years of professional experience do you have?",
        0, 30, 5,
        key="years_experience_slider"
    )
    
    responses["current_role"] = st.text_input(
        "What is your current role or most recent position?",
        "",
        key="current_role_input"
    )
    
    responses["background_summary"] = st.text_area(
        "Briefly summarize your professional background (fields, roles, industries)",
        "",
        key="background_summary_area"
    )
    
    # Skills and expertise
    st.write("### Skills & Expertise")
    responses["technical_skills"] = st.text_area(
        "What technical skills do you have? List your strongest skills first.",
        "",
        key="technical_skills_area"
    )
    
    responses["soft_skills"] = st.text_area(
        "What soft skills or personal strengths do you possess?",
        "",
        key="soft_skills_area"
    )
    
    # Work preferences
    st.write("### Work Style & Preferences")
    
    responses["work_environment"] = st.radio(
        "Do you prefer working:",
        ["Independently", "In small teams", "In large collaborative groups"],
        key="work_environment_radio"
    )
    
    responses["management_preference"] = st.selectbox(
        "Do you prefer being an individual contributor or a people manager?",
        ["Individual Contributor", "People Manager", "Combination/Depends", "Not sure yet"],
        key="management_preference_select"
    )
    
    responses["risk_tolerance"] = st.slider(
        "How comfortable are you with risk in your career? (1 = Very risk-averse, 10 = Very risk-tolerant)",
        1, 10, 5,
        key="risk_tolerance_slider"
    )
    
    # Goals and motivations
    st.write("### Goals & Motivations")
    
    responses["career_goals"] = st.text_area(
        "What are your key career goals for the next 3-5 years?",
        "",
        key="career_goals_area"
    )
    
    responses["motivations"] = st.multiselect(
        "What motivates you most in your career? (Select up to 3)",
        [
            "Financial success", "Work-life balance", "Creative expression",
            "Technical mastery", "Leadership", "Making an impact", "Job security",
            "Status/recognition", "Building something new", "Intellectual challenge",
            "Autonomy/independence", "Other"
        ],
        [],
        key="motivations_multiselect"
    )
    
    responses["timeframe"] = st.selectbox(
        "What is your preferred timeframe for significant career advancement?",
        ["0-2 years (short-term)", "3-5 years (medium-term)", "5+ years (long-term)"],
        key="timeframe_select"
    )
    
    # Challenges and constraints
    st.write("### Challenges & Constraints")
    
    responses["biggest_challenges"] = st.text_area(
        "What do you see as your biggest challenges or obstacles in your career path?",
        "",
        key="biggest_challenges_area"
    )
    
    responses["geographic_constraints"] = st.radio(
        "Do you have geographic constraints for your career?",
        ["Fully remote only", "Specific location required", "Flexible/willing to relocate", "No constraints"],
        key="geographic_constraints_radio"
    )
    
    responses["time_investment"] = st.slider(
        "How many hours per week can you realistically invest in career development activities?",
        1, 30, 10,
        key="time_investment_slider"
    )
    
    # Additional information
    st.write("### Additional Information")
    
    responses["learning_style"] = st.multiselect(
        "How do you prefer to learn? (Select all that apply)",
        [
            "Reading books/articles", "Online courses", "Hands-on projects",
            "Formal education", "Mentorship", "Social learning/communities",
            "Video tutorials", "Learning on the job"
        ],
        [],
        key="learning_style_multiselect"
    )
    
    responses["additional_info"] = st.text_area(
        "Is there anything else you'd like to share to help personalize your roadmap?",
        "",
        key="additional_info_area"
    )
    
    return responses

def ai_roadmap_generator_page(pathway=None, pathways_df=None, metrics_data=None):
    """
    Streamlit page for the AI-powered roadmap generator.
    
    Args:
        pathway (dict, optional): A pre-selected pathway
        pathways_df (DataFrame): The pathways dataframe
        metrics_data (dict): Information about metrics
    """
    st.title("Personalized Career Roadmap: Your Path to Success")
    st.write("""
    Input your resume, background information, and dream job to get a personalized career plan. 
    Our AI will create a tailored roadmap that bridges your current skills with your career aspirations.
    """)
    
    # Dream job input section
    st.write("## Step 1: Define Your Career Aspirations")
    
    # Initialize dream_job in session state if not present
    if 'dream_job' not in st.session_state:
        st.session_state.dream_job = ""
    
    # Input for dream job with description
    st.write("Enter a specific role title or describe your ideal position:")
    dream_job = st.text_input("Dream Job", 
                             value=st.session_state.dream_job, 
                             placeholder="e.g., Senior Product Manager at a tech startup",
                             key="dream_job_input")
    
    # Save dream job to session state for persistence
    st.session_state.dream_job = dream_job
    
    # Alternative: Select from existing pathways
    st.write("## Step 2: Select a Career Pathway")
    
    use_dream_job = False
    if dream_job:
        use_dream_job = st.checkbox("Use my dream job description for roadmap generation", value=True)
    
    if not use_dream_job:
        if pathway is None and pathways_df is not None:
            # Let user select a pathway
            pathway_options = pathways_df["name"].tolist()
            selected_pathway_name = st.selectbox("Select a career pathway:", pathway_options)
            
            # Get the full pathway data
            pathway = pathways_df[pathways_df["name"] == selected_pathway_name].iloc[0].to_dict()
        elif pathway is not None:
            st.success(f"Using pre-selected pathway: {pathway['name']}")
        else:
            st.error("Please select a pathway to generate a roadmap.")
            return
    else:
        # Create a temporary pathway based on the dream job
        if pathway is None:
            pathway = {
                "name": f"Custom: {dream_job}",
                "description": f"Custom career path based on your dream job: {dream_job}",
                "category": "Custom",
                "metrics": {}
            }
            st.success(f"Using your dream job: {dream_job}")
    
    # Create tabs for different input methods
    input_tab1, input_tab2, input_tab3 = st.tabs(["Questionnaire", "Upload Resume", "Combined Approach"])
    
    with input_tab1:
        st.write("## Step 2: Complete the Career Questionnaire")
        questionnaire_responses = create_questionnaire()
        
        generate_button1 = st.button("Generate Roadmap from Questionnaire", key="btn_questionnaire")
        
        if generate_button1:
            with st.spinner("Analyzing your responses and generating your personalized roadmap..."):
                # Process the questionnaire responses
                user_analysis = process_questionnaire(questionnaire_responses)
                
                # Generate the roadmap
                roadmap = generate_ai_roadmap(pathway, {
                    "questionnaire_analysis": user_analysis,
                    "questionnaire_responses": questionnaire_responses
                }, metrics_data)
                
                # Display the roadmap
                display_ai_roadmap(roadmap)
    
    with input_tab2:
        st.write("## Step 3: Upload Your Resume")
        st.write("""
        Upload your resume to have the AI analyze your background and create a roadmap 
        tailored to your experience and skills. This helps identify the gaps between your current 
        abilities and your dream job requirements.
        """)
        
        uploaded_file = st.file_uploader("Upload your resume (PDF, DOCX, or TXT)", 
                                        type=["pdf", "docx", "txt"],
                                        key="resume_upload")
                                        
        st.markdown("### Professional Background")
        current_role = st.text_input("Current Role", placeholder="e.g., Software Engineer at IBM")
        years_experience = st.slider("Years of Professional Experience", 0, 30, 3)
        
        # Skills section
        st.markdown("### Key Skills")
        st.write("List your top skills (technical and soft skills)")
        skills_input = st.text_area("Enter your skills, separated by commas", 
                                   placeholder="e.g., Python, project management, data analysis, leadership")
        
        if uploaded_file is not None:
            # Read the file content
            file_content = uploaded_file.read()
            
            # For PDFs and DOCXs, try to extract text (this is simplified)
            if uploaded_file.type == "application/pdf":
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                    temp_file.write(file_content)
                    temp_file_path = temp_file.name
                
                try:
                    import PyPDF2
                    with open(temp_file_path, 'rb') as pdf_file:
                        pdf_reader = PyPDF2.PdfReader(pdf_file)
                        file_content = ""
                        for page in pdf_reader.pages:
                            file_content += page.extract_text()
                except:
                    st.error("Unable to process PDF. Please try a text (.txt) file instead.")
                    file_content = ""
                
                # Clean up the temp file
                os.remove(temp_file_path)
            
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                    temp_file.write(file_content)
                    temp_file_path = temp_file.name
                
                try:
                    import docx
                    doc = docx.Document(temp_file_path)
                    file_content = "\n".join([para.text for para in doc.paragraphs])
                except:
                    st.error("Unable to process DOCX. Please try a text (.txt) file instead.")
                    file_content = ""
                
                # Clean up the temp file
                os.remove(temp_file_path)
            
            else:  # Text file
                file_content = file_content.decode('utf-8')
            
            st.write("### Resume Preview")
            preview = file_content[:1000] + "..." if len(file_content) > 1000 else file_content
            st.text_area("First 1000 characters of your resume:", preview, height=200, key="resume_preview")
            
            generate_button2 = st.button("Generate Roadmap from Resume", key="btn_resume")
            
            if generate_button2:
                with st.spinner("Analyzing your resume and generating your personalized roadmap..."):
                    # Process the resume
                    resume_analysis = process_resume(file_content)
                    
                    # Generate the roadmap
                    roadmap = generate_ai_roadmap(pathway, {
                        "resume_analysis": resume_analysis,
                        "resume_content": file_content[:5000]  # Limit to first 5000 chars for API
                    }, metrics_data)
                    
                    # Display the roadmap
                    display_ai_roadmap(roadmap)
    
    with input_tab3:
        st.write("## Step 2: Combined Approach (Recommended)")
        st.write("""
        For the most personalized roadmap, both upload your resume and complete the questionnaire.
        This gives the AI the most complete picture of your background and preferences.
        """)
        
        # Resume upload for combined approach
        uploaded_file_combined = st.file_uploader("Upload your resume (PDF, DOCX, or TXT)", 
                                                 type=["pdf", "docx", "txt"],
                                                 key="combined_upload")
        
        resume_content = None
        if uploaded_file_combined is not None:
            # Similar file processing as above
            file_content = uploaded_file_combined.read()
            
            if uploaded_file_combined.type == "application/pdf":
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                    temp_file.write(file_content)
                    temp_file_path = temp_file.name
                
                try:
                    import PyPDF2
                    with open(temp_file_path, 'rb') as pdf_file:
                        pdf_reader = PyPDF2.PdfReader(pdf_file)
                        resume_content = ""
                        for page in pdf_reader.pages:
                            resume_content += page.extract_text()
                except:
                    st.error("Unable to process PDF. Please try a text (.txt) file instead.")
                    resume_content = ""
                
                # Clean up the temp file
                os.remove(temp_file_path)
            
            elif uploaded_file_combined.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                    temp_file.write(file_content)
                    temp_file_path = temp_file.name
                
                try:
                    import docx
                    doc = docx.Document(temp_file_path)
                    resume_content = "\n".join([para.text for para in doc.paragraphs])
                except:
                    st.error("Unable to process DOCX. Please try a text (.txt) file instead.")
                    resume_content = ""
                
                # Clean up the temp file
                os.remove(temp_file_path)
            
            else:  # Text file
                resume_content = file_content.decode('utf-8')
                
            # Add resume preview
            st.write("### Resume Preview")
            preview = resume_content[:1000] + "..." if len(resume_content) > 1000 else resume_content
            st.text_area("First 1000 characters of your resume:", preview, height=200, key="combined_preview")
        
        # Questionnaire for combined approach with unique keys
        st.write("### Career Questionnaire")
        
        # Create a copy of the create_questionnaire function with unique keys for combined approach
        def create_combined_questionnaire():
            # Initialize responses dictionary
            responses = {}
            
            # Background section
            st.write("#### Background")
            responses["years_experience"] = st.slider(
                "How many years of professional experience do you have?",
                0, 30, 5,
                key="combined_years_experience"
            )
            
            responses["current_role"] = st.text_input(
                "What is your current role or most recent position?",
                "",
                key="combined_current_role"
            )
            
            responses["background_summary"] = st.text_area(
                "Briefly summarize your professional background (fields, roles, industries)",
                "",
                key="combined_background_summary"
            )
            
            # Skills and expertise
            st.write("#### Skills & Expertise")
            responses["technical_skills"] = st.text_area(
                "What technical skills do you have? List your strongest skills first.",
                "",
                key="combined_technical_skills"
            )
            
            responses["soft_skills"] = st.text_area(
                "What soft skills or personal strengths do you possess?",
                "",
                key="combined_soft_skills"
            )
            
            # Work preferences
            st.write("#### Work Style & Preferences")
            
            responses["work_environment"] = st.radio(
                "Do you prefer working:",
                ["Independently", "In small teams", "In large collaborative groups"],
                key="combined_work_environment"
            )
            
            responses["management_preference"] = st.selectbox(
                "Do you prefer being an individual contributor or a people manager?",
                ["Individual Contributor", "People Manager", "Combination/Depends", "Not sure yet"],
                key="combined_management_preference"
            )
            
            responses["risk_tolerance"] = st.slider(
                "How comfortable are you with risk in your career? (1 = Very risk-averse, 10 = Very risk-tolerant)",
                1, 10, 5,
                key="combined_risk_tolerance"
            )
            
            # Goals and motivations
            st.write("#### Goals & Motivations")
            
            responses["career_goals"] = st.text_area(
                "What are your key career goals for the next 3-5 years?",
                "",
                key="combined_career_goals"
            )
            
            responses["motivations"] = st.multiselect(
                "What motivates you most in your career? (Select up to 3)",
                [
                    "Financial success", "Work-life balance", "Creative expression",
                    "Technical mastery", "Leadership", "Making an impact", "Job security",
                    "Status/recognition", "Building something new", "Intellectual challenge",
                    "Autonomy/independence", "Other"
                ],
                [],
                key="combined_motivations"
            )
            
            responses["timeframe"] = st.selectbox(
                "What is your preferred timeframe for significant career advancement?",
                ["0-2 years (short-term)", "3-5 years (medium-term)", "5+ years (long-term)"],
                key="combined_timeframe"
            )
            
            # Challenges and constraints
            st.write("#### Challenges & Constraints")
            
            responses["biggest_challenges"] = st.text_area(
                "What do you see as your biggest challenges or obstacles in your career path?",
                "",
                key="combined_biggest_challenges"
            )
            
            responses["geographic_constraints"] = st.radio(
                "Do you have geographic constraints for your career?",
                ["Fully remote only", "Specific location required", "Flexible/willing to relocate", "No constraints"],
                key="combined_geographic_constraints"
            )
            
            responses["time_investment"] = st.slider(
                "How many hours per week can you realistically invest in career development activities?",
                1, 30, 10,
                key="combined_time_investment"
            )
            
            # Additional information
            st.write("#### Additional Information")
            
            responses["learning_style"] = st.multiselect(
                "How do you prefer to learn? (Select all that apply)",
                [
                    "Reading books/articles", "Online courses", "Hands-on projects",
                    "Formal education", "Mentorship", "Social learning/communities",
                    "Video tutorials", "Learning on the job"
                ],
                [],
                key="combined_learning_style"
            )
            
            responses["additional_info"] = st.text_area(
                "Is there anything else you'd like to share to help personalize your roadmap?",
                "",
                key="combined_additional_info"
            )
            
            return responses
        
        # Get questionnaire responses with unique keys
        questionnaire_responses_combined = create_combined_questionnaire()
        
        if resume_content is not None:
            generate_button3 = st.button("Generate Comprehensive Roadmap", key="btn_combined")
            
            if generate_button3:
                with st.spinner("Analyzing all your information and generating your comprehensive personalized roadmap..."):
                    # Process both resume and questionnaire
                    resume_analysis = process_resume(resume_content)
                    user_analysis = process_questionnaire(questionnaire_responses_combined)
                    
                    # Generate the roadmap with combined data
                    roadmap = generate_ai_roadmap(pathway, {
                        "resume_analysis": resume_analysis,
                        "questionnaire_analysis": user_analysis,
                        "questionnaire_responses": questionnaire_responses_combined,
                        "resume_content": resume_content[:5000]  # Limit to first 5000 chars for API
                    }, metrics_data)
                    
                    # Display the roadmap
                    display_ai_roadmap(roadmap)
        else:
            st.info("Please upload your resume to use the combined approach.")