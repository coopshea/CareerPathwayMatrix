import streamlit as st
import pandas as pd
import numpy as np
import time
import os
import json
from data import load_data, get_pathway_details, get_metrics_info
from visualizations import create_matrix_visualization
from recommendations import calculate_pathway_matches
from roadmaps import roadmap_generator_page
from ai_roadmap import ai_roadmap_generator_page
from job_postings_merged import job_posting_page
import tempfile
from skills_analysis import skills_analysis_page
from skill_graph import skill_graph_page
from utils import create_pathway_card, DEFAULT_IMAGES
from streamlit_chat import message
from user_auth import UserAuth
from database import Session, User, ChatMessage, UserDocument

# Configure page
st.set_page_config(
    page_title="CareerPath Navigator",
    page_icon="🧭",
    layout="wide"
)

# Load the data
pathways_data, metrics_data, categories = load_data()

# Main header
st.markdown("""
    <div style='text-align: center'>
        <h1>CareerPath Navigator</h1>
        <h3>Find your path, build your skills, achieve your career goals</h3>
    </div>
""", unsafe_allow_html=True)

# Initialize session state variables
if "user" not in st.session_state:
    st.session_state.user = None

# MANDATORY LOGIN CHECK
# If user is not logged in, show only the login form and nothing else
if st.session_state.user is None:
    st.markdown("### Sign in to access CareerPath Navigator")
    st.markdown("Please enter your information to proceed. This helps us understand who is using the application.")
    
    # Create two columns for the form and info
    login_col1, login_col2 = st.columns([3, 2])
    
    with login_col1:
        # Create the sign-in form
        with st.form("login_form", clear_on_submit=False):
            full_name = st.text_input("Full Name", key="name_input", placeholder="Your full name")
            email = st.text_input("Email", key="email_input", placeholder="your.email@example.com")
            organization = st.text_input("Organization (Optional)", key="org_input", placeholder="Company or school name")
            
            submitted = st.form_submit_button("Sign In")
            
            if submitted:
                # Simple validation
                if not full_name or not email or "@" not in email:
                    st.error("Please enter a valid name and email address to continue.")
                else:
                    # Create a simple hash for the user identifier
                    import hashlib
                    user_id = hashlib.md5(email.encode()).hexdigest()
                    
                    # Store user information
                    user_info = {
                        "id": user_id,
                        "name": full_name,
                        "email": email,
                        "organization": organization,
                        "created_at": time.time(),
                        "last_login": time.time()
                    }
                    
                    # Save to session state
                    st.session_state.user = user_info
                    
                    # Save to database (to be implemented in the future)
                    try:
                        session = Session()
                        db_user = session.query(User).filter(User.id == user_id).first()
                        if db_user:
                            # Update existing user
                            db_user.name = full_name
                            db_user.email = email
                            db_user.last_login = float(time.time())
                        else:
                            # Create new user
                            db_user = User(
                                id=user_id,
                                name=full_name,
                                email=email,
                                created_at=float(time.time()),
                                last_login=float(time.time())
                            )
                            session.add(db_user)
                        session.commit()
                        session.close()
                    except Exception as e:
                        print(f"Error saving user to database: {e}")
                    
                    st.success(f"Welcome, {full_name}!")
                    st.rerun()
    
    with login_col2:
        st.markdown("""
        ### Why Sign In?
        
        Your information helps us:
        
        - Track who is using the platform
        - Contact you about updates and features
        - Save your career preferences
        - Provide personalized recommendations
        - Remember your uploads and chat history
        
        We respect your privacy and do not share your data with third parties.
        """)
    
    # Stop execution here to force login
    st.stop()

# If we get here, the user is logged in
# Show small user info and logout in the corner
auth_col1, auth_col2 = st.columns([3, 1])
with auth_col2:
    st.write(f"**Signed in as:** {st.session_state.user['name']}")
    if st.button("Sign Out", key="nav_signout_btn"):
        st.session_state.user = None
        st.rerun()

# Initialize chat messages if not already in session state
if "messages" not in st.session_state:
    # Check if user is logged in to personalize greeting
    greeting_name = f", {st.session_state.user['name'].split()[0]}" if st.session_state.user else ""
    st.session_state.messages = [
        {"role": "assistant", "content": f"Hello{greeting_name}! I'm your AI career assistant. To get the most out of CareerPath Navigator, I recommend:\n\n1. Fill out the career preferences questionnaire in the 'Find Your Pathway' tab\n2. Upload your resume in the 'Skill Graph' tab for skill analysis\n3. Return here for personalized career guidance based on your profile\n\nHow can I help you today?"}
    ]

# AI Chat Assistant Helper for the welcome page
def ai_chat_assistant():
    st.write("### AI Career Assistant")
    
    # Show disclaimer
    st.markdown("""
    ℹ️ **Disclaimer**
    
    This AI assistant can provide guidance about CareerPath Navigator features and career advice.
    It has a limit of 10 messages per session to ensure fair usage.
    """)
    
    # Add caveat below the chat assistant
    st.warning("**FUNCTIONAL PROTOTYPE** - This application demonstrates core functionality (AI integration, skills extraction, etc). UI/UX design is not finalized.")
    
    # Initialize message counter
    if 'message_count' not in st.session_state:
        st.session_state.message_count = 0
        st.session_state.last_reset_time = time.time()
    
    # Initialize reflective questions sequence
    if "reflective_questions" not in st.session_state:
        st.session_state.reflective_questions = [
            "What are your biggest strengths professionally? Think about what others consistently praise you for.",
            "Which skills do you most enjoy using in your work or projects?",
            "What aspects of your current or past roles have felt most meaningful to you?",
            "Imagine your ideal workday - what activities would you be doing?",
            "What values are most important to you in your work environment?",
            "What obstacles seem to consistently appear in your career path?",
            "If resources and time weren't factors, what career path would you pursue?",
            "What specific impact do you want to make through your work?"
        ]
        st.session_state.current_reflective_q = 0
        st.session_state.reflective_mode = False
        st.session_state.reflective_responses = {}
    
    # Check if 30 minutes have passed since the last reset
    current_time = time.time()
    if current_time - st.session_state.last_reset_time > 1800:  # 1800 seconds = 30 minutes
        st.session_state.message_count = 0
        st.session_state.last_reset_time = current_time
    
    # Display remaining messages counter
    remaining_messages = max(0, 10 - st.session_state.message_count)
    if remaining_messages < 3:
        st.warning(f"**Rate limit:** {remaining_messages}/10 messages left for this session. Resets in {30 - int((current_time - st.session_state.last_reset_time)/60)} minutes.")
    else:
        st.info(f"**Messages remaining:** {remaining_messages}/10 for this session. Resets every 30 minutes.")
    
    # Option to start reflective journey
    if not st.session_state.reflective_mode:
        if st.button("Start Career Reflection Journey", key="start_reflection"):
            st.session_state.reflective_mode = True
            # Add the first reflection question as an assistant message
            first_question = st.session_state.reflective_questions[0]
            if 'messages' not in st.session_state:
                st.session_state.messages = []
            st.session_state.messages.append({"role": "assistant", "content": f"**Career Reflection:** {first_question}"})
            st.rerun()
    
    # Display chat messages from history on app rerun
    if 'messages' in st.session_state:
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
    else:
        st.session_state.messages = []
    
    # Accept user input
    if prompt := st.chat_input("Share your thoughts or ask a question..."):
        # Check rate limit
        if st.session_state.message_count >= 10:
            with st.chat_message("assistant"):
                st.error("You've reached the maximum number of messages for this session. Please wait for the rate limit to reset.")
            return
        
        # Add user message to chat history
        st.session_state.messages.append({"role": "user", "content": prompt})
        
        # Display user message in chat message container
        with st.chat_message("user"):
            st.markdown(prompt)
        
        # Increment message counter
        st.session_state.message_count += 1
        
        # Handle the response differently based on mode
        if st.session_state.reflective_mode:
            # Save the response to the session state
            try:
                # Get the current question
                current_q = st.session_state.reflective_questions[st.session_state.current_reflective_q]
                
                # Store the response
                st.session_state.reflective_responses[current_q] = prompt
                
                # Move to the next question
                st.session_state.current_reflective_q += 1
                
                # Create a placeholder for the assistant's response
                with st.chat_message("assistant"):
                    response_placeholder = st.empty()
                    
                    # If there are more questions, ask the next one
                    if st.session_state.current_reflective_q < len(st.session_state.reflective_questions):
                        next_q = st.session_state.reflective_questions[st.session_state.current_reflective_q]
                        response = f"Thank you for sharing that. Let's continue with another reflection: **{next_q}**"
                    else:
                        # End of reflective questions
                        response = "Thank you for completing the career reflection! I've recorded your responses, which will help personalize your career guidance. Is there anything specific you'd like to explore now based on your reflections?"
                        st.session_state.reflective_mode = False
                        
                        # In a full implementation, save all responses to database
                        from database import ChatMessage
                        # Logic to save to database would go here
                        
                    # Display the response
                    response_placeholder.markdown(response)
                
                # Add assistant response to chat history
                st.session_state.messages.append({"role": "assistant", "content": response})
                
            except Exception as e:
                with st.chat_message("assistant"):
                    response = f"I appreciate your response. There was a technical issue, but we can continue our conversation. What else would you like to discuss about your career journey?"
                    st.markdown(response)
                st.session_state.messages.append({"role": "assistant", "content": response})
                print(f"Error in reflective mode: {e}")
        else:
            # Normal chat mode - process the user's question
            with st.chat_message("assistant"):
                response_placeholder = st.empty()
                
                # Stream the response from API
                full_response = stream_ai_response(prompt, response_placeholder)
                
            # Add assistant response to chat history
            st.session_state.messages.append({"role": "assistant", "content": full_response})
    
    # Add controls for reflective mode
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("Clear Chat", key="clear_chat"):
            st.session_state.messages = []
            st.session_state.current_reflective_q = 0
            st.session_state.reflective_mode = False
            st.session_state.reflective_responses = {}
            st.rerun()
    
    with col2:
        if not st.session_state.reflective_mode and st.button("New Reflection", key="new_reflection"):
            st.session_state.current_reflective_q = 0
            st.session_state.reflective_mode = True
            # Add the first reflection question
            first_question = st.session_state.reflective_questions[0]
            st.session_state.messages.append({"role": "assistant", "content": f"**Career Reflection:** {first_question}"})
            st.rerun()

# Stream AI-powered response with word-by-word display
def stream_ai_response(question, placeholder):
    """Generate and stream an AI-powered response using OpenAI API"""
    # Handle None case
    if question is None:
        default_response = "I can help you explore different career paths and develop your skills. What would you like to know?"
        placeholder.markdown(default_response)
        return default_response
        
    try:
        # Get the API key from environment variable
        api_key = os.environ.get("OPENAI_API_KEY")
        if not api_key:
            fallback = "I couldn't access the OpenAI API key. Using rule-based responses instead.\n\n" + get_quick_response(question)
            placeholder.markdown(fallback)
            return fallback
        
        from openai import OpenAI
        
        # Create OpenAI client
        client = OpenAI(api_key=api_key)
        
        # System message for the chat context
        system_message = """You are a helpful career assistant in the CareerPath Navigator application. 
        You should encourage users to:
        1. Fill out the career preferences questionnaire in the 'Find Your Pathway' tab
        2. Upload their resume in the 'Skill Graph' tab for skill analysis
        3. Return to chat for personalized guidance based on their profile
        
        You can guide users to different features:
        - 2x2 Matrix (tab 1): For comparing career paths visually
        - Find Your Pathway (tab 2): For matching preferences to careers
        - Basic Roadmap (tab 3): For generating simple career roadmaps
        - AI Roadmap (tab 4): For generating AI-powered personalized roadmaps
        - Job Posting (tab 5): For analyzing job opportunities
        - Skills Analysis (tab 6): For finding high-impact skills in the market
        - Skill Graph (tab 7): For analyzing user skills and gaps
        
        Keep responses friendly, concise and helpful. Always recommend the appropriate tool
        based on what the user is trying to accomplish."""
        
        # Generate streaming response with context about the application
        stream = client.chat.completions.create(
            model="gpt-4o",  # Use the latest model
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": question}
            ],
            max_tokens=250,
            stream=True
        )
        
        # Process the streaming response
        collected_chunks = []
        collected_content = ""
        
        # Display the streaming response in real-time
        for chunk in stream:
            if chunk.choices[0].delta.content is not None:
                content_piece = chunk.choices[0].delta.content
                collected_chunks.append(content_piece)
                collected_content += content_piece
                placeholder.markdown(collected_content + "▌")
        
        # Display the final response without cursor
        placeholder.markdown(collected_content)
        
        # Extract any tab guidance from the AI response
        tab_keywords = {
            "2x2 matrix": 1,
            "career matrix": 1, 
            "find your pathway": 2,
            "matching": 2,
            "preferences": 2,
            "basic roadmap": 3,
            "ai roadmap": 4,
            "roadmap generator": 4,
            "job posting": 5,
            "analyze job": 5,
            "skills analysis": 6,
            "high-impact skills": 6,
            "skill graph": 7,
            "skill gaps": 7
        }
        
        # Check if any keywords appear in the response
        if collected_content and isinstance(collected_content, str):
            response_lower = collected_content.lower()
            for keyword, tab_idx in tab_keywords.items():
                if keyword in response_lower:
                    st.session_state.active_tab = tab_idx
                    break
                
        return collected_content
        
    except Exception as e:
        # Fall back to rule-based response
        fallback = f"I encountered an error with the AI response. Using rule-based responses instead.\n\n" + get_quick_response(question)
        placeholder.markdown(fallback)
        return fallback

# Get AI-powered response using OpenAI (non-streaming version)
def get_ai_response(question):
    """Generate an AI-powered response using OpenAI API (for backward compatibility)"""
    # Handle None case
    if question is None:
        return "I can help you explore different career paths and develop your skills. What would you like to know?"
        
    try:
        # Get the API key from environment variable
        api_key = os.environ.get("OPENAI_API_KEY")
        if not api_key:
            return "I couldn't access the OpenAI API key. Using rule-based responses instead.\n\n" + get_quick_response(question)
        
        from openai import OpenAI
        
        # Create OpenAI client
        client = OpenAI(api_key=api_key)
        
        # Generate response with context about the application
        ai_response = client.chat.completions.create(
            model="gpt-4o",  # Use the latest model
            messages=[
                {"role": "system", "content": """You are a helpful career assistant in the CareerPath Navigator application. 
                You should encourage users to:
                1. Fill out the career preferences questionnaire in the 'Find Your Pathway' tab
                2. Upload their resume in the 'Skill Graph' tab for skill analysis
                3. Return to chat for personalized guidance based on their profile
                
                You can guide users to different features:
                - 2x2 Matrix (tab 1): For comparing career paths visually
                - Find Your Pathway (tab 2): For matching preferences to careers
                - Basic Roadmap (tab 3): For generating simple career roadmaps
                - AI Roadmap (tab 4): For generating AI-powered personalized roadmaps
                - Job Posting (tab 5): For analyzing job opportunities
                - Skills Analysis (tab 6): For finding high-impact skills in the market
                - Skill Graph (tab 7): For analyzing user skills and gaps
                
                Keep responses friendly, concise and helpful. Always recommend the appropriate tool
                based on what the user is trying to accomplish."""},
                {"role": "user", "content": question}
            ],
            max_tokens=250
        )
        
        # Get the response content
        response = ai_response.choices[0].message.content
        
        # Extract any tab guidance from the AI response
        tab_keywords = {
            "2x2 matrix": 1,
            "career matrix": 1, 
            "find your pathway": 2,
            "matching": 2,
            "preferences": 2,
            "basic roadmap": 3,
            "ai roadmap": 4,
            "roadmap generator": 4,
            "job posting": 5,
            "analyze job": 5,
            "skills analysis": 6,
            "high-impact skills": 6,
            "skill graph": 7,
            "skill gaps": 7
        }
        
        # Check if any keywords appear in the response
        if response is not None and isinstance(response, str):
            response_lower = response.lower()
            for keyword, tab_idx in tab_keywords.items():
                if keyword in response_lower:
                    st.session_state.active_tab = tab_idx
                    break
                
        return response
        
    except Exception as e:
        # Fall back to rule-based response
        return f"I encountered an error with the AI response. Using rule-based responses instead.\n\n" + get_quick_response(question)

# Provide quick rule-based responses based on keywords for low latency
def get_quick_response(question):
    """Provide quick responses based on keywords for low latency"""
    # Default to empty string if question is None
    if question is None:
        return "I can help you explore different career paths and develop your skills. What would you like to know?"
        
    question = question.lower()
    
    # Career guidance prompt responses
    if "resume" in question and ("upload" in question or "analyze" in question):
        st.session_state.active_tab = 7  # Skill Graph tab
        return "I recommend uploading your resume in the Skill Graph tab. This will help us analyze your skills and identify gaps. I've selected that tab for you!"
        
    if "questionnaire" in question or "preferences" in question:
        st.session_state.active_tab = 2  # Find Your Pathway tab
        return "To find career paths that match your preferences, please fill out the questionnaire in the Find Your Pathway tab. I've selected that tab for you!"
        
    if "personalized" in question or "custom" in question or "for me" in question:
        return "For truly personalized guidance, I recommend:\n\n1. Complete the questionnaire in the 'Find Your Pathway' tab\n2. Upload your resume in the 'Skill Graph' tab\n\nOnce you've done both, I can provide much more tailored advice!"
    
    # Simple mapping of keywords to responses and tab navigation
    response_map = {
        "compare": (1, "I'll show you the Career Matrix to compare different paths."),
        "career path": (1, "Let's explore career paths in the 2x2 Matrix."),
        "explore": (1, "The Career Matrix is perfect for exploration."),
        "matrix": (1, "Opening the 2x2 Matrix visualization for you."),
        
        "skill": (7, "Let's check out your skills in the Skill Graph."),
        "gap": (7, "The Skill Graph will help identify your skill gaps."),
        "analyze": (7, "I'll help you analyze your skills with our Skill Graph."),
        
        "job": (5, "Let's analyze a job posting for you."),
        "posting": (5, "The Job Posting analyzer will help evaluate opportunities."),
        "opportunity": (5, "I'll open the Job Posting analyzer for you."),
        
        "roadmap": (4, "Let's create a career roadmap for you."),
        "plan": (4, "The AI Roadmap generator will help you plan your career."),
        "develop": (4, "Let's develop your career plan with the AI Roadmap."),
        
        "recommend": (2, "Let's find pathways that match your preferences."),
        "match": (2, "The Pathway Finder will help match you with careers."),
        
        "market": (6, "Let's analyze market trends in Skills Analysis."),
        "demand": (6, "The Skills Analysis tool shows high-demand skills."),
        "high impact": (6, "I'll show you high-impact skills in the market.")
    }
    
    # Check for keyword matches
    for keyword, (tab_index, response) in response_map.items():
        if keyword in question:
            st.session_state.active_tab = tab_index
            return response
    
    # Default response if no match
    return "I can help you explore career paths, analyze skills, and create personalized roadmaps. For the best experience, I recommend filling out the questionnaire in the 'Find Your Pathway' tab and uploading your resume in the 'Skill Graph' tab. What would you like to do today?"

# Create tabs
# Initialize active tab if not already set
if 'active_tab' not in st.session_state:
    st.session_state.active_tab = 0

# Define tab names in the correct order
tab_names = [
    "Welcome", "Job & Resume Analysis", "AI Roadmap", "Skill Graph", 
    "Project Portfolio", "2x2 Matrix", "Find Your Pathway", "About"
]

# Create the tabs at the top of the page
tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8 = st.tabs(tab_names)

# Welcome tab content
with tab1:
    # Button grid - 3x2 layout
    st.write("")  # Add some space at the top
    
    # Create 3 rows of 2 buttons each
    row1 = st.columns(2)
    with row1[0]:
        if st.button("📊 Compare Career Paths", use_container_width=True, key="btn_matrix"):
            st.session_state.active_tab = 1  # 2x2 Matrix tab
            st.rerun()
    
    with row1[1]:
        if st.button("💡 Find Matching Careers", use_container_width=True, key="btn_pathway"):
            st.session_state.active_tab = 2  # Find Your Pathway tab
            st.rerun()
    
    row2 = st.columns(2)
    with row2[0]:
        if st.button("🧩 Analyze My Skills", use_container_width=True, key="btn_skills"):
            st.session_state.active_tab = 7  # Skill Graph tab
            st.rerun()
    
    with row2[1]:
        if st.button("🛣️ Create Career Roadmap", use_container_width=True, key="btn_roadmap"):
            st.session_state.active_tab = 4  # AI Roadmap tab
            st.rerun()
    
    row3 = st.columns(2)
    with row3[0]:
        if st.button("🔍 Analyze Job Posting", use_container_width=True, key="btn_job"):
            st.session_state.active_tab = 5  # Job Posting tab
            st.rerun()
    
    with row3[1]:
        if st.button("📈 Find High-Impact Skills", use_container_width=True, key="btn_market"):
            st.session_state.active_tab = 6  # Skills Analysis tab
            st.rerun()
        
        # AI chat assistant below the buttons
        st.markdown("---")
        st.markdown("### Not sure where to start?")
        
        # AI chat assistant
        ai_chat_assistant()

# 2x2 Matrix tab
with tab2:
        st.write("## Find Your Ideal Career Path")
        st.write("""
        Visualize career options based on how well they match your skills and interests. The 2x2 matrix helps you identify
        paths that align with your strengths and passions, making career decisions easier and more fulfilling.
        """)
        
        # Only show sidebar in this tab
        with st.sidebar:
            st.markdown("## Matrix Controls")
            
            # Category filters - add key parameter to fix duplicate element issue
            selected_categories = st.multiselect(
                "Filter by Category",
                ["All"] + categories,
                default=["All"],
                key="matrix_categories"
            )
            
            if "All" in selected_categories:
                filtered_pathways = pathways_data
            else:
                filtered_pathways = pathways_data[pathways_data['category'].isin(selected_categories)]
            
            # Select x and y axis metrics with defaults focused on skills and enjoyment
            x_metric = st.selectbox(
                "X-Axis Metric", 
                list(metrics_data.keys()),
                index=list(metrics_data.keys()).index("skill_match") if "skill_match" in metrics_data else 0,
                key="x_axis_metric"
            )
            
            y_metric = st.selectbox(
                "Y-Axis Metric", 
                list(metrics_data.keys()),
                index=list(metrics_data.keys()).index("enjoyment_factor") if "enjoyment_factor" in metrics_data else 0,
                key="y_axis_metric"
            )
            
            # Additional info for selected metrics
            st.markdown(f"### {metrics_data[x_metric]['name']}")
            st.write(metrics_data[x_metric]['description'])
            st.markdown(f"### {metrics_data[y_metric]['name']}")
            st.write(metrics_data[y_metric]['description'])
        
        # Create the matrix visualization
        figure = create_matrix_visualization(filtered_pathways, x_metric, y_metric, metrics_data)
        st.plotly_chart(figure, use_container_width=True)
        
        # Display selected pathway details
        st.write("## Pathway Details")
        if 'selected_pathway' not in st.session_state:
            st.session_state.selected_pathway = None
        
        # Instructions if no pathway is selected
        if st.session_state.selected_pathway is None:
            st.info("👆 Click on a pathway in the matrix above to see detailed information")
        else:
            # Get details for the selected pathway
            pathway_details = get_pathway_details(pathways_data, st.session_state.selected_pathway)
            
            # Create columns for the pathway details
            col1, col2 = st.columns([2, 3])
            
            with col1:
                st.write(f"### {pathway_details['name']}")
                st.write(f"**Category:** {pathway_details['category']}")
                st.write(f"**Description:** {pathway_details['description']}")
                st.write("**Target Customers:** " + pathway_details['target_customers'])
                st.write("**Success Examples:** " + ", ".join(pathway_details['success_examples']))
                st.write("**Key Skills:**")
                for skill in pathway_details['key_skills']:
                    st.markdown(f"- {skill}")
            
            with col2:
                # Create expandable sections for pathway details
                with st.expander("Key Metrics", expanded=True):
                    # Show key metrics in a clean table format
                    key_metrics = ["risk_level", "success_probability", "terminal_value", "expected_value_10yr"]
                    key_data = []
                    
                    for metric in key_metrics:
                        if metric in pathway_details['metrics']:
                            metric_info = pathway_details['metrics'][metric]
                            key_data.append({
                                "Metric": metrics_data[metric]['name'],
                                "Value": f"{metric_info['value']}/10",
                                "Interpretation": metrics_data[metric]['rubric'][str(metric_info['value'])]
                            })
                    
                    st.table(pd.DataFrame(key_data))
                    
                with st.expander("All Metrics", expanded=False):
                    # Show all metrics in a table format
                    all_data = []
                    for metric_key, metric_info in pathway_details['metrics'].items():
                        all_data.append({
                            "Metric": metrics_data[metric_key]['name'],
                            "Value": f"{metric_info['value']}/10",
                            "Interpretation": metrics_data[metric_key]['rubric'][str(metric_info['value'])]
                        })
                    
                    st.table(pd.DataFrame(all_data))
                    
                with st.expander("Sources & Evidence", expanded=False):
                    selected_metric = st.selectbox(
                        "Select metric to view sources",
                        list(pathway_details['metrics'].keys()),
                        format_func=lambda x: metrics_data[x]['name'],
                        key="source_metric"
                    )
                    
                    if 'sources' in pathway_details['metrics'][selected_metric]:
                        for source in pathway_details['metrics'][selected_metric]['sources']:
                            st.write(f"**{source['title']}** - {source['publisher']}")
                            st.write(f"> {source['extract']}")
                            st.write(f"[Read more]({source['url']})")
                    else:
                        st.write("No sources available for this metric.")
                    
                with st.expander("Rationale", expanded=False):
                    if 'rationale' in pathway_details:
                        for metric, explanation in pathway_details['rationale'].items():
                            st.write(f"**{metrics_data[metric]['name']}:** {explanation}")
                    else:
                        st.write("No detailed rationale available for this pathway.")
                
                # Add buttons to generate roadmaps for this pathway
                roadmap_cols = st.columns(2)
                with roadmap_cols[0]:
                    if st.button("Generate Basic Roadmap", key="gen_basic"):
                        # Set a session state variable to indicate we want to switch to the basic roadmap tab
                        # with this pathway pre-selected
                        st.session_state.generate_roadmap_for = pathway_details
                        # Switch to the basic roadmap tab (index 3)
                        st.session_state.active_tab = 3
                        st.rerun()
                with roadmap_cols[1]:
                    if st.button("Generate AI-Powered Roadmap", key="gen_ai"):
                        # Set a session state variable to indicate we want to switch to the AI roadmap tab
                        # with this pathway pre-selected
                        st.session_state.generate_roadmap_for = pathway_details
                        # Switch to the AI roadmap tab (index 4)
                        st.session_state.active_tab = 4
                        st.rerun()

# Find Your Pathway tab
with tab3:
        st.write("## Find Pathways That Match Your Preferences")
        st.write("""
        Use the sliders below to indicate your preferences for different aspects of a career pathway. 
        The tool will calculate how well each pathway matches your preferences and show the best matches.
        """)
        
        # First show reflective questions to help users think about their preferences
        with st.expander("Reflective Questions to Consider", expanded=False):
            questions = [
                {
                    "topic": "Risk Tolerance",
                    "question": "Think about past career decisions. How did you feel when taking risks versus pursuing safer options? Did you feel more energized by stability or by potential upside?",
                    "relevance": "Understanding your personal risk tolerance helps identify pathways that match your comfort level."
                },
                {
                    "topic": "Technical Specialization",
                    "question": "Consider times when you've deeply specialized versus remained a generalist. Which approach felt more natural to you?",
                    "relevance": "Some pathways require deep technical expertise while others reward breadth of knowledge."
                },
                {
                    "topic": "Time Horizon",
                    "question": "Reflect on your patience with delayed gratification. Are you willing to invest years before seeing significant returns?",
                    "relevance": "Pathways differ dramatically in how quickly they deliver financial returns."
                },
                {
                    "topic": "Control vs. Optionality",
                    "question": "Consider situations where you had to choose between maintaining full control and preserving future options. Which did you prioritize?",
                    "relevance": "Certain paths offer greater autonomy while others maintain flexibility to pivot."
                },
                {
                    "topic": "Network Reliance",
                    "question": "How comfortable are you with success depending on your ability to build and leverage professional relationships?",
                    "relevance": "Pathways vary significantly in their dependency on networks and connections."
                }
            ]
            
            for q_idx, q in enumerate(questions):
                st.write(f"**{q['topic']}**")
                st.write(q['question'])
                st.write(f"*{q['relevance']}*")
                st.write("---")
        
        # Create columns for preference input
        col1, col2 = st.columns(2)
        
        # Dictionary to store user preferences
        user_preferences = {}
        importance_weights = {}
        
        # Select which metrics to show in the preference input
        preference_metrics = ["risk_level", "capital_requirements", "technical_specialization", 
                             "network_dependency", "time_to_return", "control", "optionality"]
        
        # Create preference sliders in two columns
        for i, metric in enumerate(preference_metrics):
            metric_info = metrics_data[metric]
            col = col1 if i % 2 == 0 else col2
            
            with col:
                st.write(f"### {metric_info['name']}")
                st.write(metric_info['description'])
                
                # Create preference range slider with unique key
                min_val, max_val = st.slider(
                    f"Your preferred range for {metric_info['name']}",
                    1, 10, (3, 7),
                    help=f"Low: {metric_info['low_label']} | High: {metric_info['high_label']}",
                    key=f"pref_range_{metric}"
                )
                
                # Store preference in dictionary
                user_preferences[metric] = (min_val, max_val)
                
                # Create importance slider with unique key
                importance = st.slider(
                    f"How important is {metric_info['name']} to you?",
                    1, 10, 5,
                    help="1 = Not important | 10 = Extremely important",
                    key=f"importance_{metric}"
                )
                
                # Store importance in dictionary
                importance_weights[metric] = importance
                
                # Show interpretation of slider values
                st.write(f"*Low (1): {metric_info['rubric']['1']}*")
                st.write(f"*High (10): {metric_info['rubric']['10']}*")
                st.write("---")
        
        # Calculate pathway matches based on user preferences
        st.write("## Your Recommended Pathways")
        
        if st.button("Find Matching Pathways", key="find_match"):
            matches = calculate_pathway_matches(pathways_data, user_preferences, importance_weights)
            
            # Display the top 5 matches
            for i, (pathway_id, match_score, match_explanation) in enumerate(matches[:5]):
                pathway = pathways_data[pathways_data['id'] == pathway_id].iloc[0]
                
                st.markdown(f"### {i+1}. {pathway['name']} - {match_score:.0f}% Match")
                
                # Create two columns for the recommendation display
                rec_col1, rec_col2 = st.columns([1, 2])
                
                with rec_col1:
                    # Display the match explanation
                    st.write("**Why this matches your preferences:**")
                    for metric, explanation in match_explanation.items():
                        st.write(f"- {metrics_data[metric]['name']}: {explanation}")
                    
                    # Buttons for details and roadmap generation
                    if st.button(f"See Full Details #{i+1}", key=f"details_{i}"):
                        st.session_state.selected_pathway = pathway_id
                        st.session_state.active_tab = 1  # Switch to 2x2 Matrix tab
                        st.rerun()
                    
                    if st.button(f"Generate AI Roadmap #{i+1}", key=f"ai_roadmap_{i}"):
                        pathway_details = get_pathway_details(pathways_data, pathway_id)
                        st.session_state.generate_roadmap_for = pathway_details
                        st.session_state.active_tab = 4  # Switch to AI Roadmap tab
                        st.rerun()
                
                with rec_col2:
                    # Create a small visualization of how this pathway matches user preferences
                    matches_data = []
                    for metric in preference_metrics:
                        if metric in pathway['metrics']:
                            min_pref, max_pref = user_preferences[metric]
                            pathway_value = pathway['metrics'][metric]['value']
                            in_range = min_pref <= pathway_value <= max_pref
                            
                            matches_data.append({
                                "Metric": metrics_data[metric]['name'],
                                "Your Preference": f"{min_pref}-{max_pref}",
                                "Pathway Value": pathway_value,
                                "Match": "✅" if in_range else "❌"
                            })
                    
                    st.table(pd.DataFrame(matches_data))
                    
                st.write("---")

# Project Portfolio tab (replacing Basic Roadmap)
with tab6:
        st.title("Project Portfolio")
        st.write("""
        Document your projects and skills to build a comprehensive portfolio that demonstrates your abilities.
        This helps potential employers see concrete evidence of your capabilities.
        """)
        
        # Upload project documentation
        st.header("Add Project to Portfolio")
        
        # Project details
        project_name = st.text_input("Project Name", placeholder="e.g., E-commerce Website Redesign")
        project_description = st.text_area("Project Description", placeholder="Describe the project, your role, and the impact it had")
        
        # Project document upload
        st.write("Upload documentation or examples of your project (optional)")
        project_doc = st.file_uploader("Upload project document or image", type=["pdf", "docx", "txt", "jpg", "png"])
        
        # Skills utilized
        st.subheader("Skills Demonstrated")
        st.write("Select the skills you practiced or demonstrated in this project")
        
        # Get skills from database
        job_skills = []
        try:
            from database import fetch_job_skills
            job_skills = fetch_job_skills(top_n=30)  # Get top 30 skills
            skill_names = [skill["name"] for skill in job_skills]
        except Exception as e:
            st.error(f"Could not fetch skills from database: {e}")
            skill_names = ["Project Management", "Data Analysis", "Web Development", "Content Creation", 
                          "UX Design", "Marketing", "Leadership", "Python", "JavaScript"]
        
        # Multi-select skills with rating
        selected_skills = st.multiselect("Skills used in this project", skill_names)
        
        skill_ratings = {}
        if selected_skills:
            st.write("Rate your proficiency with each skill used in this project:")
            for skill in selected_skills:
                skill_ratings[skill] = st.slider(f"{skill} proficiency", 1, 5, 3, key=f"skill_{skill}")
        
        # Add project button
        if st.button("Add Project to Portfolio"):
            if project_name and project_description:
                try:
                    # Save project data (simplified version for demo)
                    st.success(f"Project '{project_name}' added to your portfolio!")
                    st.write("In a production version, this would be saved to your user profile in the database.")
                    
                    if project_doc:
                        st.write(f"Document '{project_doc.name}' attached to your project.")
                    
                    if selected_skills:
                        st.write("Skills demonstrated:")
                        for skill in selected_skills:
                            st.write(f"- {skill}: Proficiency level {skill_ratings.get(skill, 3)}/5")
                except Exception as e:
                    st.error(f"Error saving project: {e}")
            else:
                st.warning("Please provide at least a project name and description.")

# AI Roadmap tab - now tab3
with tab3:
        # Check if we're coming from a pathway detail view with a pre-selected pathway
        pre_selected_pathway = None
        if 'generate_roadmap_for' in st.session_state:
            pre_selected_pathway = st.session_state.generate_roadmap_for
            # Clear the session state so it doesn't persist
            del st.session_state.generate_roadmap_for
            
        ai_roadmap_generator_page(pre_selected_pathway, pathways_data, metrics_data)

# Job & Resume Analysis tab - now tab2 (improved organization)
with tab2:
        st.title("Job & Resume Analysis")
        st.write("""
        Upload your resume and job descriptions to analyze the alignment between your skills and job requirements.
        This analysis helps identify which skills are worth learning first and provides valuable context for personalized career guidance.
        """)
        
        # Create tabs for file upload and manual entry
        file_tab, questionnaire_tab = st.tabs(["Upload Documents (Recommended)", "Fill Questionnaire"])
        
        with file_tab:
            # Create columns for side-by-side layout
            col1, col2 = st.columns(2)
            
            with col1:
                st.header("Resume Analysis")
                st.write("Upload your resume to extract and analyze your skills")
                
                # Resume upload
                resume_file = st.file_uploader("Upload your resume (PDF, DOCX, or TXT)", 
                                              type=["pdf", "docx", "txt"],
                                              key="resume_upload_combined")
                
                if resume_file:
                    # Process resume
                    file_content = resume_file.read()
                    
                    # Extract text based on file type
                    resume_text = ""
                    try:
                        if resume_file.type == "application/pdf":
                            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                                temp_file.write(file_content)
                                temp_file_path = temp_file.name
                            
                            # Extract text from PDF
                            import PyPDF2
                            with open(temp_file_path, 'rb') as pdf_file:
                                pdf_reader = PyPDF2.PdfReader(pdf_file)
                                for page in pdf_reader.pages:
                                    resume_text += page.extract_text()
                            
                            # Clean up
                            os.remove(temp_file_path)
                        
                        elif resume_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                                temp_file.write(file_content)
                                temp_file_path = temp_file.name
                            
                            # Extract text from DOCX
                            import docx
                            doc = docx.Document(temp_file_path)
                            resume_text = "\n".join([para.text for para in doc.paragraphs])
                            
                            # Clean up
                            os.remove(temp_file_path)
                        
                        else:  # Text file
                            resume_text = file_content.decode('utf-8')
                        
                        # Display resume text preview
                        st.write("### Resume Preview")
                        preview = resume_text[:300] + "..." if len(resume_text) > 300 else resume_text
                        st.text_area("Preview:", preview, height=100)
                        
                        # Process resume with AI button
                        process_resume_btn = st.button("Analyze Resume")
                    
                    except Exception as e:
                        st.error(f"Error processing file: {e}")
                
                else:
                    st.info("Upload your resume to extract and analyze your skills")
                    process_resume_btn = False
            
            with col2:
                st.header("Job Posting Analysis")
                st.write("Upload a job description to analyze its requirements")
                
                # Job posting analysis (using file upload as primary method)
                uploaded_file = st.file_uploader("Upload a job posting document", 
                                              type=["pdf", "docx", "txt"],
                                              key="job_file_merged")
                
                job_text = None
                
                if uploaded_file:
                    # Process the file based on type
                    try:
                        if uploaded_file.type == "application/pdf":
                            import PyPDF2
                            # Save the uploaded file temporarily
                            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                                temp_file.write(uploaded_file.read())
                                temp_file_path = temp_file.name
                            
                            # Extract text from the PDF
                            with open(temp_file_path, 'rb') as file:
                                pdf_reader = PyPDF2.PdfReader(file)
                                job_text = ""
                                for page_num in range(len(pdf_reader.pages)):
                                    job_text += pdf_reader.pages[page_num].extract_text()
                            
                            # Clean up the temporary file
                            os.remove(temp_file_path)
                            
                        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                            import docx
                            # Save the uploaded file temporarily
                            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                                temp_file.write(uploaded_file.read())
                                temp_file_path = temp_file.name
                            
                            # Extract text from the DOCX
                            doc = docx.Document(temp_file_path)
                            job_text = "\n".join([para.text for para in doc.paragraphs])
                            
                            # Clean up the temporary file
                            os.remove(temp_file_path)
                            
                        else:  # Text file
                            job_text = uploaded_file.read().decode('utf-8')
                        
                        # Display job posting preview
                        st.write("### Job Posting Preview")
                        preview = job_text[:300] + "..." if len(job_text) > 300 else job_text
                        st.text_area("Preview:", preview, height=100)
                        
                        # Save to session state
                        st.session_state.job_posting_text = job_text
                        
                        # Process job posting button
                        process_job_btn = st.button("Analyze Job Posting")
                    
                    except Exception as e:
                        st.error(f"Error processing file: {e}")
                        process_job_btn = False
                
                else:
                    st.info("Upload a job description to analyze requirements")
                    st.write("Or paste text instead:")
                    job_text = st.text_area("Paste job posting text", height=100, key="job_text_merged")
                    process_job_btn = st.button("Analyze Job Text") if job_text else False
                    if job_text:
                        st.session_state.job_posting_text = job_text
            
            # Process the files if buttons were clicked
            if 'resume_text' in locals() and process_resume_btn:
                with st.spinner("Analyzing your resume..."):
                    from skill_graph import analyze_resume_skills
                    
                    # Extract skills from resume
                    try:
                        skills = analyze_resume_skills(resume_text)
                        
                        # Display extracted skills
                        st.write("### Extracted Skills")
                        skill_col1, skill_col2 = st.columns(2)
                        
                        for i, (skill, info) in enumerate(list(skills.items())[:6]):
                            rating = info.get('rating', 3)
                            context = info.get('context', '')
                            
                            # Create a visual representation of rating
                            rating_stars = "★" * rating + "☆" * (5 - rating)
                            
                            with skill_col1 if i % 2 == 0 else skill_col2:
                                st.write(f"**{skill}** - {rating_stars}")
                                if context:
                                    st.write(f"*{context}*")
                                st.write("---")
                        
                        # Save to session state for the AI chat to access
                        st.session_state.user_resume_skills = skills
                        
                    except Exception as e:
                        st.error(f"Error analyzing resume: {e}")
                        # Sample skills data
                        skills = {
                            "Python": {"rating": 4, "context": "5 years experience with data analysis"},
                            "Project Management": {"rating": 3, "context": "Led team of 5 developers"},
                            "Communication": {"rating": 5, "context": "Customer-facing role"},
                        }
                        st.write("### Sample Skills (Demo)")
                        for skill, info in skills.items():
                            rating = info.get('rating', 3)
                            context = info.get('context', '')
                            rating_stars = "★" * rating + "☆" * (5 - rating)
                            st.write(f"**{skill}** - {rating_stars}")
                            if context:
                                st.write(f"*{context}*")
                            st.write("---")
            
            if job_text and process_job_btn:
                with st.spinner("Analyzing job posting..."):
                    try:
                        # Import necessary function
                        from job_postings_merged import analyze_job_posting
                        
                        # Call the function to analyze the job posting
                        analysis = analyze_job_posting(job_text)
                        
                        # Display summary
                        st.write("### Job Analysis Summary")
                        job_col1, job_col2 = st.columns(2)
                        
                        with job_col1:
                            st.write(f"**Job Title:** {analysis.get('title', 'Unknown')}")
                            st.write(f"**Company:** {analysis.get('company', 'Unknown')}")
                            st.write(f"**Category:** {analysis.get('category', 'Unknown')}")
                            
                            st.write("**Required Skills:**")
                            for skill in analysis.get('required_skills', [])[:5]:  # Limit to top 5
                                st.write(f"- {skill}")
                            if len(analysis.get('required_skills', [])) > 5:
                                st.write(f"- *and {len(analysis.get('required_skills', [])) - 5} more*")
                        
                        with job_col2:
                            st.write(f"**Level:** {analysis.get('seniority', 'Not specified')}")
                            st.write(f"**Location:** {analysis.get('location', 'Not specified')}")
                            
                            if analysis.get('preferred_skills'):
                                st.write("**Preferred Skills:**")
                                for skill in analysis.get('preferred_skills', [])[:5]:  # Limit to top 5
                                    st.write(f"- {skill}")
                                if len(analysis.get('preferred_skills', [])) > 5:
                                    st.write(f"- *and {len(analysis.get('preferred_skills', [])) - 5} more*")
                        
                        # Store in session state for resume skill comparison
                        st.session_state.job_skills = {
                            'required': analysis.get('required_skills', []),
                            'preferred': analysis.get('preferred_skills', [])
                        }
                        
                        # Save the analysis to session state
                        st.session_state.job_analysis = analysis
                        
                    except Exception as e:
                        st.error(f"Error analyzing job posting: {e}")
            
            # Display combined analysis if both resume and job are analyzed
            if 'user_resume_skills' in st.session_state and 'job_skills' in st.session_state:
                st.success("✅ Both resume and job posting analyzed! View the skill gap analysis below.")
                
                st.write("## Skill Gap Analysis")
                
                # Get required skills from job
                required_skills = set(st.session_state.job_skills.get('required', []))
                preferred_skills = set(st.session_state.job_skills.get('preferred', []))
                
                # Get user's skills
                user_skills = set(st.session_state.user_resume_skills.keys())
                
                # Calculate matches and gaps
                required_matches = required_skills.intersection(user_skills)
                required_gaps = required_skills.difference(user_skills)
                preferred_matches = preferred_skills.intersection(user_skills)
                preferred_gaps = preferred_skills.difference(user_skills)
                
                # Calculate match percentage for required skills
                if required_skills:
                    match_percentage = len(required_matches) / len(required_skills) * 100
                else:
                    match_percentage = 0
                
                # Display match percentage
                st.markdown(f"### Overall Match: {match_percentage:.1f}%")
                
                # Display skills columns
                col1, col2 = st.columns(2)
                
                with col1:
                    st.write("#### 🎯 Skills You Already Have")
                    if required_matches:
                        st.write("**Required skills you have:**")
                        for skill in required_matches:
                            st.write(f"✓ {skill}")
                    
                    if preferred_matches:
                        st.write("**Preferred skills you have:**")
                        for skill in preferred_matches:
                            st.write(f"✓ {skill}")
                    
                    if not required_matches and not preferred_matches:
                        st.info("No direct skill matches found. Consider developing the skills listed on the right.")
                
                with col2:
                    st.write("#### 📝 Skills to Develop")
                    if required_gaps:
                        st.write("**Required skills to learn:**")
                        for skill in required_gaps:
                            st.write(f"→ {skill}")
                    
                    if preferred_gaps:
                        st.write("**Preferred skills to consider:**")
                        for skill in preferred_gaps:
                            st.write(f"→ {skill}")
                    
                    if not required_gaps and not preferred_gaps:
                        st.success("Congratulations! You have all the skills listed in the job posting.")
                
                # Next steps based on analysis
                st.write("### Recommended Next Steps")
                st.info("""
                1. Focus on developing the required skills you're missing
                2. Generate an AI roadmap for this career path
                3. Add the skills you already have to your portfolio with evidence
                """)
                
                if st.button("Generate AI Roadmap for This Job"):
                    # Set the dream job based on the analyzed job posting
                    job_title = st.session_state.job_analysis.get('title', 'Unknown Position')
                    st.session_state.dream_job = job_title
                    # Go to AI Roadmap tab
                    st.session_state.active_tab = 2  # Index of AI Roadmap tab
                    st.rerun()
        
        with questionnaire_tab:
            st.header("Career Preferences Questionnaire")
            st.write("If you don't have a resume or job posting, fill out this questionnaire to get career recommendations.")
            
            # Dream job input
            st.write("#### Your Career Aspirations")
            dream_job = st.text_input("What's your dream job or role?", key="dream_job_questionnaire")
            
            # Current skills
            st.write("#### Your Current Skills")
            current_skills = st.text_area("List your key skills (comma separated)", placeholder="e.g., Python, Project Management, Data Analysis", key="skills_questionnaire")
            
            # Work preferences
            st.write("#### Work Preferences")
            work_pref_options = ["Remote work", "Flexible hours", "Creative freedom", "Structured environment", 
                                "Team collaboration", "Independent work", "Fast-paced", "Work-life balance"]
            work_preferences = st.multiselect("Select your work preferences", work_pref_options, key="work_pref_questionnaire")
            
            # Industry interests
            st.write("#### Industry Interests")
            industry_options = ["Technology", "Healthcare", "Finance", "Education", "Entertainment", 
                               "Manufacturing", "Retail", "Government", "Non-profit"]
            industries = st.multiselect("Select industries you're interested in", industry_options, key="industries_questionnaire")
            
            # Save and analyze button
            if st.button("Save and Analyze Preferences"):
                st.session_state.questionnaire_responses = {
                    "dream_job": dream_job,
                    "current_skills": [skill.strip() for skill in current_skills.split(",") if skill.strip()],
                    "work_preferences": work_preferences,
                    "industries": industries
                }
                
                st.success("Preferences saved! Now you can explore career pathways based on your inputs.")
                
                # Go to Find Your Pathway tab
                if st.button("View Recommended Pathways"):
                    st.session_state.active_tab = 6  # Index of Find Your Pathway tab
                    st.rerun()
        
        # Benefits of analysis at the bottom
        st.subheader("Benefits of Analysis")
        st.write("""
        By analyzing your background and job requirements, the CareerPath Navigator can:
        - Identify which skills are worth learning first
        - Score career opportunities for the 2x2 matrix
        - Give the AI chatbot personal context about your background
        - Create a personalized learning roadmap targeted at your dream job
        """)

# Skill Graph tab - now tab4
with tab4:
    # Load the skill graph page
    skill_graph_page()

# 2x2 Matrix tab - now tab6 
with tab6:
    st.title("Career Pathways Matrix")
    st.write("""
    This 2x2 matrix visualizes different career pathways based on two important metrics:
    skill alignment (how well your skills match the requirements) and growth potential.
    """)
    
    # Get metrics for the matrix axes
    x_metric = "skill_match"
    y_metric = "growth_potential"
    
    # Create the visualization
    create_matrix_visualization(pathways_data, x_metric, y_metric, metrics_data)

# Find Your Pathway tab - now tab7
with tab7:
    st.title("Find Your Pathway")
    st.write("""
    Complete the questionnaire below to discover career pathways that align with your preferences.
    """)
    
    # Create the questionnaire
    st.header("Career Preferences Questionnaire")
    
    metrics = metrics_data.keys()
    user_preferences = {}
    importance_weights = {}
    
    for metric_id in metrics:
        metric_info = metrics_data[metric_id]
        
        st.write(f"### {metric_info['name']}")
        st.write(metric_info['description'])
        
        # Ask for preference within a range
        preference_range = st.slider(
            f"Where do you prefer to be on the {metric_info['name']} scale?",
            0, 100, (30, 70),
            key=f"pref_{metric_id}"
        )
        
        # Ask for importance
        importance = st.slider(
            f"How important is {metric_info['name']} to you?",
            0, 10, 5,
            key=f"imp_{metric_id}"
        )
        
        user_preferences[metric_id] = preference_range
        importance_weights[metric_id] = importance
    
    # Button to calculate matches
    if st.button("Find Matching Pathways"):
        # Calculate matches
        matches = calculate_pathway_matches(pathways_data, user_preferences, importance_weights)
        
        # Display matches
        st.write("## Your Top Pathway Matches")
        
        # Display top 3 matches
        for i, (pathway_id, score, explanation) in enumerate(matches[:3]):
            pathway = get_pathway_details(pathways_data, pathway_id)
            
            st.write(f"### {i+1}. {pathway['name']} ({score:.1f}% match)")
            st.write(pathway['description'])
            
            with st.expander("Why this pathway matches your preferences"):
                st.write(explanation)
            
            # Button to view detailed pathway
            if st.button(f"View Detailed Pathway", key=f"view_{i}"):
                st.session_state.selected_pathway_id = pathway_id
                st.session_state.active_tab = 6  # Now the 2x2 Matrix tab
                st.rerun()
            
            st.write("---")
        
        # Display remaining matches in a table
        if len(matches) > 3:
            st.write("### Other Potential Matches")
            
            # Prepare data for the table
            matches_data = []
            for pathway_id, score, _ in matches[3:10]:  # Show up to 10 total
                pathway = get_pathway_details(pathways_data, pathway_id)
                matches_data.append({
                    "Pathway": pathway['name'],
                    "Match Score": f"{score:.1f}%",
                    "Category": pathway['category']
                })
            
            # Display as a table
            st.table(pd.DataFrame(matches_data))

# About tab - now tab8
with tab8:
    # About page content
    st.image(DEFAULT_IMAGES["data_viz_concept"], use_container_width=True)
    
    st.write("## About CareerPath Navigator")
    
    st.write("""CareerPath Navigator is a comprehensive career development platform that leverages AI 
to transform skill management, career planning, and professional growth through intelligent, 
data-driven insights.

### Our Mission

To help professionals at any stage understand themselves better, identify fitting career paths, 
and create actionable plans to bridge the gap between their current skills and future goals.

### Key Features

- **Visual Career Exploration**: Compare career options using an interactive 2x2 matrix
- **AI-Powered Analysis**: Extract skills from resumes and analyze job postings
- **Skill Gap Identification**: See which high-demand skills you should develop
- **Personalized Roadmaps**: Get customized step-by-step guidance
- **Project Portfolio**: Document your work to demonstrate your skills

### Technologies

Built with Streamlit, PostgreSQL, and OpenAI integration.
""")
    
    # Add database health status indicator
    with st.expander("System Status", expanded=False):
        from database import test_connection, check_migration_needed
        
        connection_ok = test_connection()
        if connection_ok:
            st.success("✅ Database connection is working properly.")
        else:
            st.error("❌ Database connection failed. Using file-based data instead.")
            
        migration_needed = check_migration_needed()
        if migration_needed:
            st.info("Database schema will be automatically updated on next startup.")
        else:
            st.success("✅ Database schema is up to date.")