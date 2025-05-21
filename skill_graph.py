import streamlit as st
import networkx as nx
from pyvis.network import Network
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import os
import json
import tempfile
from database import fetch_job_skills, save_user_skills, fetch_user_skills, JobSkill, Session
from io import BytesIO
import base64
import random
import re
from datetime import datetime

# Define AI skills analysis function
def analyze_resume_skills(resume_text):
    """
    Analyze resume text to extract skills using OpenAI.
    
    Args:
        resume_text (str): The text content of the resume
        
    Returns:
        dict: Dictionary of skills with ratings and context
    """
    try:
        # Import OpenAI client
        from openai import OpenAI
        
        # Initialize client with API key from environment variable
        client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
        
        # Create the prompt
        prompt = f"""
        You are a skilled resume analyzer that extracts and evaluates professional skills.
        
        Analyze the following resume text and extract a comprehensive list of professional skills:
        
        {resume_text}
        
        For each skill:
        1. Assign a rating from 1-5 based on apparent proficiency
        2. Include a brief context about how they've used this skill
        
        Format your response as a valid JSON object with skills as keys and objects containing 'rating' and 'experience' as values.
        Example: 
        {{
            "Python": {{
                "rating": 4,
                "experience": "5 years of experience with data analysis and web development"
            }}
        }}
        
        Only include hard skills, soft skills, and technical competencies that are clearly evident from the resume.
        """
        
        # Make API call to OpenAI
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You extract and rate professional skills from resumes."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2,
            response_format={"type": "json_object"}
        )
        
        # Parse the response content as JSON
        skills_data = json.loads(response.choices[0].message.content)
        return skills_data
        
    except Exception as e:
        st.error(f"Error analyzing resume: {str(e)}")
        return get_sample_skills()


def get_sample_skills():
    """Return sample skills when API extraction fails"""
    return {
        "Python": {
            "rating": 4,
            "experience": "5 years of experience with data analysis and automation"
        },
        "Project Management": {
            "rating": 3,
            "experience": "Led multiple cross-functional teams on successful projects"
        },
        "Data Analysis": {
            "rating": 4,
            "experience": "Performed statistical analysis on large datasets"
        },
        "Communication": {
            "rating": 5,
            "experience": "Regular presenter at team meetings and industry conferences"
        },
        "JavaScript": {
            "rating": 3,
            "experience": "Built interactive data visualizations for web applications"
        }
    }


def process_resume_text(file_content, file_type):
    """
    Process the uploaded resume file and extract text based on file type.
    
    Args:
        file_content: The binary content of the file
        file_type: The MIME type or extension of the file
        
    Returns:
        str: The extracted text content
    """
    try:
        # Process PDF
        if file_type == "application/pdf":
            from PyPDF2 import PdfReader
            pdf_reader = PdfReader(BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
            
        # Process DOCX
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            from docx import Document
            doc = Document(BytesIO(file_content))
            return "\n".join([para.text for para in doc.paragraphs])
            
        # Process plain text
        else:
            return file_content.decode("utf-8")
            
    except Exception as e:
        st.error(f"Error processing file: {str(e)}")
        return ""


def generate_skill_graph(user_skills, job_skills):
    """
    Generate a NetworkX graph of user skills and job skills.
    
    Args:
        user_skills (dict): User's skills with ratings
        job_skills (list): Skills from job postings
        
    Returns:
        nx.Graph: The skill graph
    """
    G = nx.Graph()
    
    # Add user skills
    for skill_name, skill_data in user_skills.items():
        rating = skill_data.get('rating', 3)
        G.add_node(skill_name, 
                  type="user_skill", 
                  rating=rating,
                  title=f"{skill_name} (Your rating: {rating}/5)")
    
    # Track job skill frequencies and types
    job_skill_data = {}
    for skill in job_skills:
        name = skill.get('name', '')
        if not name:
            continue
            
        if name not in job_skill_data:
            job_skill_data[name] = {
                'frequency': 0,
                'type': skill.get('skill_type', 'general')
            }
        job_skill_data[name]['frequency'] += 1
    
    # Add job skills
    for skill_name, data in job_skill_data.items():
        if skill_name not in G:  # Only add if not already in graph (as user skill)
            G.add_node(skill_name, 
                      type="job_skill", 
                      frequency=data['frequency'],
                      skill_type=data['type'],
                      title=f"{skill_name} (In {data['frequency']} job postings)")
    
    # Add edges - connect related skills
    # For this example, we'll just add some plausible skill relationships
    skill_relationships = {
        "Python": ["Data Analysis", "Machine Learning", "Django", "Flask", "Automation"],
        "JavaScript": ["React", "Node.js", "Web Development", "Frontend"],
        "SQL": ["Database", "Data Analysis", "PostgreSQL", "MySQL"],
        "Project Management": ["Agile", "Scrum", "Leadership", "Communication"],
        "Data Analysis": ["Statistics", "Python", "R", "Excel", "Visualization"]
    }
    
    # Connect skills based on relationships
    nodes = list(G.nodes())
    for node in nodes:
        # Try to find related skills
        related = []
        for key, values in skill_relationships.items():
            if node.lower() == key.lower():
                related = values
            elif node.lower() in [v.lower() for v in values]:
                related = [key] + [v for v in values if v.lower() != node.lower()]
        
        # Add edges for related skills
        for rel in related:
            for other_node in nodes:
                if other_node.lower() == rel.lower():
                    G.add_edge(node, other_node)
    
    # Make sure all nodes have at least one connection
    # Connect to most relevant job skills
    for node in nodes:
        if G.degree(node) == 0:
            # Find a relevant connection
            if node in user_skills:
                # Connect to a job skill
                for job_node in [n for n in nodes if n not in user_skills]:
                    G.add_edge(node, job_node)
                    break
            else:
                # Connect to a user skill
                for user_node in [n for n in nodes if n in user_skills]:
                    G.add_edge(node, user_node)
                    break
    
    return G


def create_interactive_graph(graph):
    """
    Create an interactive PyVis Network from a NetworkX graph.
    
    Args:
        graph (nx.Graph): The NetworkX graph
        
    Returns:
        pyvis.network.Network: The interactive network
    """
    # Create a PyVis network
    net = Network(height="600px", width="100%", bgcolor="#ffffff", font_color="black")
    
    # Configure physics
    net.barnes_hut(spring_length=200, spring_strength=0.05, damping=0.09)
    
    # Add nodes with properties
    for node in graph.nodes(data=True):
        node_id = node[0]
        node_attrs = node[1]
        
        # Determine color and size based on type and metrics
        if node_attrs.get('type') == 'user_skill':
            rating = node_attrs.get('rating', 3)
            color = f"rgba(0, 100, 255, {0.5 + rating/10})"  # Brighter blue for higher ratings
            size = 20 + rating * 3  # Larger nodes for higher ratings
            group = 'user_skills'
        else:  # job_skill
            frequency = node_attrs.get('frequency', 1)
            skill_type = node_attrs.get('skill_type', 'general')
            
            # Different colors for different skill types
            if skill_type == 'required':
                color = f"rgba(255, 0, 0, {0.5 + min(frequency/20, 0.5)})"  # Red for required
            elif skill_type == 'preferred':
                color = f"rgba(255, 165, 0, {0.5 + min(frequency/20, 0.5)})"  # Orange for preferred
            else:
                color = f"rgba(0, 128, 0, {0.5 + min(frequency/20, 0.5)})"  # Green for general
                
            size = 15 + min(frequency * 2, 20)  # Larger nodes for more frequent skills
            group = 'job_skills'
        
        # Add the node with appropriate attributes
        net.add_node(node_id, 
                    title=node_attrs.get('title', node_id),
                    color=color, 
                    size=size,
                    label=node_id,
                    group=group)
    
    # Add edges
    for edge in graph.edges():
        net.add_edge(edge[0], edge[1], color="rgba(120, 120, 120, 0.5)")
    
    return net


def get_html_network(net):
    """Convert the network to html for display in Streamlit"""
    # Generate HTML file
    html = net.generate_html()
    
    # Fix PyVis CSS for Streamlit
    html = html.replace('<style type="text/css">', '<style>')
    
    # Return the HTML
    return html


def identify_skill_gaps(user_skills, job_skills):
    """
    Identify skill gaps between user skills and job requirements.
    
    Args:
        user_skills (dict): User's skills with ratings
        job_skills (list): Skills from job postings
        
    Returns:
        list: Top skills to acquire, sorted by frequency in job postings
    """
    # Get user skill names
    user_skill_names = set(skill.lower() for skill in user_skills.keys())
    
    # Process job skills
    job_skill_counts = {}
    for skill in job_skills:
        name = skill.get('name', '').lower()
        if not name:
            continue
        
        if name not in job_skill_counts:
            job_skill_counts[name] = {
                'name': skill.get('name', ''),  # Keep original case
                'frequency': 0,
                'skill_type': skill.get('skill_type', 'general')
            }
        job_skill_counts[name]['frequency'] += 1
    
    # Find skills in job postings that user doesn't have
    missing_skills = []
    for name, data in job_skill_counts.items():
        # Skip if user has this skill (case-insensitive)
        if name in user_skill_names:
            continue
            
        # Check for similar skills (to avoid "Python" vs "Python Programming" issues)
        has_similar = False
        for user_skill in user_skill_names:
            if (name in user_skill or user_skill in name) and len(min(name, user_skill)) >= 4:
                has_similar = True
                break
        
        if not has_similar:
            missing_skills.append(data)
    
    # Sort by frequency
    missing_skills.sort(key=lambda x: x['frequency'], reverse=True)
    
    return missing_skills


def generate_skill_roadmap(user_skills, target_skill, all_skills):
    """
    Generate a roadmap to acquire a target skill.
    
    Args:
        user_skills (dict): User's current skills
        target_skill (str): The skill to acquire
        all_skills (list): All skills data
        
    Returns:
        dict: A roadmap with steps
    """
    # Define skill prerequisites (simplified example)
    skill_prerequisites = {
        "Machine Learning": ["Python", "Statistics", "Mathematics"],
        "Data Science": ["Python", "Statistics", "Data Analysis"],
        "React": ["JavaScript", "HTML", "CSS"],
        "Django": ["Python", "Web Development"],
        "Product Management": ["Communication", "Planning", "Risk Management"]
    }
    
    # For demonstration, check if skill exists in our map
    if target_skill not in skill_prerequisites:
        # Create a generic roadmap if skill isn't in our map
        return {
            "skill": target_skill,
            "description": f"A personalized roadmap to acquire {target_skill}",
            "prerequisites": [],
            "steps": [
                {
                    "name": f"Learn {target_skill} fundamentals",
                    "resources": [
                        "Take an online course on a platform like Coursera or Udemy",
                        "Read introductory books on the subject",
                        "Follow tutorials on YouTube"
                    ]
                },
                {
                    "name": "Practice with projects",
                    "resources": [
                        "Create small projects to apply what you've learned",
                        "Join online communities for guidance and feedback",
                        "Contribute to open source projects if applicable"
                    ]
                },
                {
                    "name": "Demonstrate your skills",
                    "resources": [
                        "Add projects to your portfolio",
                        "Create a case study showcasing your work",
                        "Take on freelance work to gain experience"
                    ]
                }
            ]
        }
    
    # Check what prerequisites the user already has
    missing_prerequisites = []
    for prereq in skill_prerequisites[target_skill]:
        if prereq not in user_skills:
            missing_prerequisites.append(prereq)
    
    # Create a roadmap with steps
    roadmap = {
        "skill": target_skill,
        "description": f"A personalized roadmap to acquire {target_skill}",
        "prerequisites": skill_prerequisites[target_skill],
        "missing_prerequisites": missing_prerequisites,
        "steps": []
    }
    
    # Add steps for missing prerequisites
    if missing_prerequisites:
        roadmap["steps"].append({
            "name": "Build prerequisite skills",
            "description": f"Before learning {target_skill}, you need to acquire these foundation skills",
            "tasks": [f"Learn {prereq}" for prereq in missing_prerequisites]
        })
    
    # Add core learning steps
    roadmap["steps"].append({
        "name": f"Learn {target_skill} fundamentals",
        "description": "Master the basic concepts and techniques",
        "tasks": [
            "Take an online course or bootcamp",
            "Read books and documentation",
            "Practice with exercises and small projects"
        ]
    })
    
    roadmap["steps"].append({
        "name": "Build practical experience",
        "description": "Apply your knowledge in real-world scenarios",
        "tasks": [
            "Create a portfolio project",
            "Join community forums and discussions",
            "Collaborate on open source projects"
        ]
    })
    
    roadmap["steps"].append({
        "name": "Demonstrate mastery",
        "description": "Show your expertise and continue improving",
        "tasks": [
            "Document your projects and process",
            "Get feedback from experienced practitioners",
            "Teach or mentor others to solidify your knowledge"
        ]
    })
    
    return roadmap


def display_roadmap(roadmap):
    """
    Display a skill acquisition roadmap in Streamlit.
    
    Args:
        roadmap (dict): The roadmap to display
    """
    st.write(f"## Roadmap to Learn {roadmap['skill']}")
    st.write(roadmap["description"])
    
    # Prerequisites section
    if roadmap["prerequisites"]:
        st.write("### Prerequisites")
        for prereq in roadmap["prerequisites"]:
            if "missing_prerequisites" in roadmap and prereq in roadmap["missing_prerequisites"]:
                st.markdown(f"- ❌ {prereq} (Need to learn)")
            else:
                st.markdown(f"- ✅ {prereq} (You already know this)")
    
    # Steps section
    st.write("### Learning Path")
    for i, step in enumerate(roadmap["steps"]):
        with st.expander(f"Step {i+1}: {step['name']}", expanded=True):
            if "description" in step:
                st.write(step["description"])
            
            if "tasks" in step:
                for task in step["tasks"]:
                    st.markdown(f"- {task}")
            
            if "resources" in step:
                st.write("**Recommended Resources:**")
                for resource in step["resources"]:
                    st.markdown(f"- {resource}")


def add_user_project(user_data, project_info):
    """
    Add a project to the user's profile.
    
    Args:
        user_data (dict): The user's profile data
        project_info (dict): Information about the project
        
    Returns:
        dict: Updated user data
    """
    # Add the project to the user's projects
    user_data["projects"].append(project_info)
    
    # Update the user's skills based on the project
    for skill in project_info["skills"]:
        # If the user already has this skill, potentially upgrade the rating
        if skill in user_data["skills"]:
            # Simple logic: increment rating if it's not already at 5
            current_rating = user_data["skills"][skill]["rating"]
            if current_rating < 5:
                user_data["skills"][skill]["rating"] = min(current_rating + 1, 5)
            
            # Add the project to the skill's project list
            if "projects" not in user_data["skills"][skill]:
                user_data["skills"][skill]["projects"] = []
            
            if project_info["name"] not in user_data["skills"][skill]["projects"]:
                user_data["skills"][skill]["projects"].append(project_info["name"])
        else:
            # Add the new skill to the user's profile
            user_data["skills"][skill] = {
                "rating": 2,  # Start with a moderate rating for new skills from projects
                "experience": f"Used in project: {project_info['name']}",
                "projects": [project_info["name"]]
            }
    
    return user_data


def analyze_project_for_skills(project_description, project_link=""):
    """
    Analyze a project to extract skills.
    In a real implementation, this would use OpenAI to analyze the project.
    
    Args:
        project_description (str): Description of the project
        project_link (str): Link to the project (optional)
        
    Returns:
        dict: Project information with extracted skills
    """
    try:
        # Import OpenAI client
        from openai import OpenAI
        
        # Initialize client with API key from environment variable
        client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
        
        # Create the prompt
        prompt = f"""
        You are a skilled project analyst that identifies professional skills demonstrated in projects.
        
        Analyze the following project description and extract a list of professional skills:
        
        Project Description: {project_description}
        Project Link: {project_link}
        
        Identify 3-7 specific skills demonstrated in this project. Focus on technical, professional, and domain-specific skills.
        """
        
        # Make API call to OpenAI
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You identify professional skills from project descriptions."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=200
        )
        
        # Extract skills from the response
        skills_text = response.choices[0].message.content
        
        # Parse skills from the text (simple approach)
        skills = []
        for line in skills_text.split('\n'):
            line = line.strip()
            if line and line.startswith(('- ', '• ', '* ')):
                skill = line[2:].strip()
                # Remove any comments in parentheses
                if '(' in skill:
                    skill = skill.split('(')[0].strip()
                skills.append(skill)
        
        # If no skills were identified in the bullet points, try to extract them directly
        if not skills:
            skills = re.findall(r'\b([A-Z][a-z]+ [A-Z][a-z]+|[A-Z][a-z]+)\b', skills_text)
            skills = list(set(skills))[:7]  # Limit to 7 unique skills
        
        # Generate a project name if not provided
        project_name = f"Project: {project_description.split()[0:3]}"
        
        # Create project info
        project_info = {
            "name": project_name,
            "description": project_description,
            "link": project_link,
            "skills": skills,
            "date_added": datetime.now().strftime("%Y-%m-%d")
        }
        
        return project_info
        
    except Exception as e:
        st.error(f"Error analyzing project: {str(e)}")
        
        # Fallback to basic skill extraction from keywords
        keywords = ["python", "javascript", "data", "design", "management", 
                   "communication", "analysis", "development", "web", "mobile"]
        
        # Simple keyword matching
        found_skills = []
        for keyword in keywords:
            if keyword.lower() in project_description.lower():
                # Capitalize first letter
                skill = keyword[0].upper() + keyword[1:]
                found_skills.append(skill)
        
        # Ensure at least some skills are returned
        if not found_skills:
            found_skills = ["Project Management", "Communication"]
        
        # Generate a project name
        words = project_description.split()
        project_name = f"Project: {' '.join(words[:3])}"
        
        return {
            "name": project_name,
            "description": project_description,
            "link": project_link,
            "skills": found_skills,
            "date_added": datetime.now().strftime("%Y-%m-%d")
        }


def render_skill_graph_tab(user_data=None, selectbox=None):
    """Render the skill graph tab with the provided user data"""
    
    # Initialize session state variables if they don't exist
    if "skills_updated" not in st.session_state:
        st.session_state.skills_updated = False
    
    # Always initialize with empty skills - this ensures each user gets their own clean slate
    if "user_skills" not in st.session_state:
        st.session_state.user_skills = {}
    
    # Initialize job_skills as empty for new users (no demo data)
    if "job_skills" not in st.session_state:
        st.session_state.job_skills = {}
    
    # Create tabs for different skill analysis views
    skills_tab, graph_tab, roadmap_tab = st.tabs(["Skills Profile", "Skill Graph", "Skill Roadmap"])
    
    with skills_tab:
        # No header, just the content
        st.write("Upload your resume and job descriptions to analyze your skills and match them with job requirements.")
        
        # Create two columns for resume and job posting analysis
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("### 📄 Upload Resume")
            st.markdown("Upload your resume to automatically extract your skills and experience.")
            resume_file = st.file_uploader("Upload your resume (PDF, DOCX, or TXT)", 
                                      type=["pdf", "docx", "txt"],
                                      key="skill_graph_resume_upload")
            
            if resume_file:
                # Process resume
                with st.spinner("Processing resume..."):
                    file_content = resume_file.read()
                    
                    # Extract text based on file type
                    resume_text = ""
                    try:
                        if resume_file.type == "application/pdf":
                            # Process PDF
                            import io
                            from PyPDF2 import PdfReader
                            pdf_reader = PdfReader(io.BytesIO(file_content))
                            for page in pdf_reader.pages:
                                resume_text += page.extract_text() + "\n"
                        
                        elif resume_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                            # Process DOCX
                            import io
                            from docx import Document
                            doc = Document(io.BytesIO(file_content))
                            resume_text = "\n".join([para.text for para in doc.paragraphs])
                        
                        else:
                            # Process as plain text
                            resume_text = file_content.decode("utf-8")
                            
                        # Show extracted text in expander
                        with st.expander("View Extracted Text", expanded=False):
                            st.text_area("Resume Content", resume_text, height=300, disabled=True)
                        
                        # Store in session state
                        if user_data:
                            user_data.resume_bytes = file_content
                        if "resume_text" not in st.session_state:
                            st.session_state.resume_text = resume_text
                        
                        # Analyze resume for skills
                        with st.spinner("Extracting skills..."):
                            try:
                                # Try using AI for skill extraction
                                skills_dict = analyze_resume_skills(resume_text)
                                st.session_state.user_skills = skills_dict
                                st.success("Skills extracted successfully!")
                                
                                # Save to database
                                try:
                                    from database import save_user_skills
                                    save_user_skills(skills_dict)
                                except Exception as e:
                                    st.warning(f"Skills analyzed but couldn't save to database: {e}")
                            except Exception as e:
                                # Fallback to sample skills on API error
                                st.warning(f"Using sample skills data due to an error. Try again later.")
                                st.session_state.user_skills = get_sample_skills()
                    
                    except Exception as e:
                        st.error(f"Error processing file: {str(e)}")
        
        with col2:
            st.markdown("### 💼 Analyze Job Posting")
            st.markdown("Upload a job description to identify required skills and compare with your profile.")
            
            # File uploader for job posting
            job_file = st.file_uploader("Upload job description (PDF, DOCX, or TXT)", 
                                       type=["pdf", "docx", "txt"],
                                       key="job_file_upload")
            
            # Text area for job description
            job_text = st.text_area("Or paste job description here", height=150, key="job_text_area")
            
            # Process job posting when file is uploaded or text is entered
            if job_file or job_text:
                with st.spinner("Analyzing job description..."):
                    try:
                        # Import analyze_job_posting function
                        from job_postings_merged import analyze_job_posting
                        
                        # Process file if provided
                        if job_file:
                            file_content = job_file.read()
                            job_text_from_file = ""
                            
                            # Extract text based on file type
                            if job_file.type == "application/pdf":
                                # Process PDF
                                import io
                                from PyPDF2 import PdfReader
                                pdf_reader = PdfReader(io.BytesIO(file_content))
                                for page in pdf_reader.pages:
                                    job_text_from_file += page.extract_text() + "\n"
                            
                            elif job_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                                # Process DOCX
                                import io
                                from docx import Document
                                doc = Document(io.BytesIO(file_content))
                                job_text_from_file = "\n".join([para.text for para in doc.paragraphs])
                            
                            else:
                                # Process as plain text
                                job_text_from_file = file_content.decode("utf-8")
                            
                            # Use file content as job text
                            job_text = job_text_from_file
                            
                            # Store in user data if available
                            if user_data:
                                user_data.job_bytes = file_content
                        
                        if job_text:
                            # Analyze job posting
                            job_data = analyze_job_posting(job_text)
                            
                            if job_data:
                                # Store in session state
                                if "job_skills" not in st.session_state:
                                    st.session_state.job_skills = {
                                        'required': job_data.get('required_skills', []),
                                        'preferred': job_data.get('preferred_skills', [])
                                    }
                                
                                # Store in session state for skills graph
                                if "job_postings" not in st.session_state:
                                    st.session_state.job_postings = []
                                
                                # Add job data to the list if it's not already there
                                if job_data not in st.session_state.job_postings:
                                    st.session_state.job_postings.append(job_data)
                                
                                # Save to database
                                try:
                                    from database import add_job_posting_to_db, add_skills_from_job
                                    from job_postings_merged import convert_job_to_pathway
                                    
                                    # Convert to pathway format
                                    pathway = convert_job_to_pathway(job_data)
                                    
                                    # Add to database
                                    add_job_posting_to_db(pathway)
                                    
                                    # Add skills to database
                                    with st.spinner("Saving skills to database..."):
                                        import sqlalchemy
                                        from sqlalchemy.orm import sessionmaker
                                        from database import engine
                                        
                                        Session = sessionmaker(bind=engine)
                                        session = Session()
                                        try:
                                            add_skills_from_job(session, pathway)
                                            session.commit()
                                        except Exception as e:
                                            session.rollback()
                                            st.warning(f"Could not save job skills to database: {e}")
                                        finally:
                                            session.close()
                                        
                                except Exception as e:
                                    st.warning(f"Job analyzed but couldn't save to database: {e}")
                                
                                st.success("Job posting analyzed successfully!")
                                
                                # Job details
                                st.subheader("Job Details")
                                st.write(f"**Position:** {job_data.get('title', 'Not specified')}")
                                st.write(f"**Company:** {job_data.get('company', 'Not specified')}")
                                st.write(f"**Category:** {job_data.get('category', 'Not specified')}")
                                
                                # Show job text in expander
                                with st.expander("View Job Text", expanded=False):
                                    st.text_area("Job Description", job_text, height=300, disabled=True)
                                
                                # Required skills
                                req_col1, req_col2 = st.columns(2)
                                
                                with req_col1:
                                    st.subheader("Required Skills")
                                    required_skills = job_data.get('required_skills', [])
                                    if required_skills:
                                        for skill in required_skills:
                                            # Check if user has this skill
                                            if "user_skills" in st.session_state and skill in st.session_state.user_skills:
                                                st.markdown(f"- ✅ {skill}")
                                            else:
                                                st.markdown(f"- ❌ {skill}")
                                    else:
                                        st.write("No required skills specified.")
                                
                                with req_col2:
                                    st.subheader("Preferred Skills")
                                    preferred_skills = job_data.get('preferred_skills', [])
                                    if preferred_skills:
                                        for skill in preferred_skills:
                                            # Check if user has this skill
                                            if "user_skills" in st.session_state and skill in st.session_state.user_skills:
                                                st.markdown(f"- ✅ {skill}")
                                            else:
                                                st.markdown(f"- ❌ {skill}")
                                    else:
                                        st.write("No preferred skills specified.")
                                
                                # Add missing skills to profile
                                if st.button("Add Missing Skills to Profile", key="add_skills_btn"):
                                    if "user_skills" not in st.session_state:
                                        st.session_state.user_skills = {}
                                    
                                    all_job_skills = set(required_skills + preferred_skills)
                                    user_skills = set(st.session_state.user_skills.keys()) if "user_skills" in st.session_state else set()
                                    missing_skills = all_job_skills - user_skills
                                    
                                    # Add each missing skill
                                    for skill in missing_skills:
                                        st.session_state.user_skills[skill] = {
                                            "rating": 1,
                                            "experience": f"Added from job: {job_data.get('title', 'Unknown')}",
                                            "projects": []
                                        }
                                    
                                    # Save to database
                                    try:
                                        from database import save_user_skills
                                        save_user_skills(st.session_state.user_skills)
                                    except Exception as e:
                                        st.warning(f"Could not save skills to database: {e}")
                                    
                                    if missing_skills:
                                        st.success(f"Added {len(missing_skills)} missing skills to your profile!")
                                    else:
                                        st.info("No new skills to add. You already have all the skills listed in this job.")
                            else:
                                st.error("Could not analyze job posting. Please try a different job description.")
                    except Exception as e:
                        st.error(f"Error analyzing job posting: {str(e)}")
                        st.info("Please check your internet connection and try again.")
        
        # Display and edit skills section
        st.subheader("Your Skills")
        
        if "user_skills" not in st.session_state or not st.session_state.user_skills:
            st.info("No skills found. Please upload a resume or job description to extract skills.")
        else:
            st.write(f"You have {len(st.session_state.user_skills)} skills in your profile:")
            
            # Create a form for adding new skills
            with st.expander("Add New Skill", expanded=False):
                with st.form("add_skill_form", clear_on_submit=True):
                    new_skill = st.text_input("Skill name:", placeholder="e.g., Python, Project Management")
                    new_rating = st.slider("Proficiency level:", 1, 5, 3)
                    new_experience = st.text_area("Experience with this skill:", placeholder="Describe your experience...", height=100)
                    
                    submitted = st.form_submit_button("Add Skill")
                    
                    if submitted and new_skill:
                        # Check if skill already exists
                        if new_skill in st.session_state.user_skills:
                            st.error(f"Skill '{new_skill}' already exists in your profile.")
                        else:
                            # Add the skill
                            st.session_state.user_skills[new_skill] = {
                                "rating": new_rating,
                                "experience": new_experience,
                                "projects": []
                            }
                            st.session_state.skills_updated = True
                            st.success(f"Added skill: {new_skill}")
                            st.rerun()
            
            # Display and edit current skills
            with st.expander("View and Edit Skills", expanded=True):
                skills_to_delete = []
                
                for skill, data in st.session_state.user_skills.items():
                    st.write(f"### {skill}")
                    
                    # Create columns for rating and removal
                    col1, col2, col3 = st.columns([2, 2, 1])
                    
                    with col1:
                        st.write(f"**{skill}**")
                    
                    with col2:
                        new_rating = st.slider(
                            f"Rate your {skill} skills", 
                            min_value=1, 
                            max_value=5, 
                            value=data["rating"],
                            key=f"rating_{skill}"
                        )
                        # Update the rating if changed
                        if new_rating != data["rating"]:
                            st.session_state.user_skills[skill]["rating"] = new_rating
                            st.session_state.skills_updated = True
                    
                    with col3:
                        if st.button("Remove", key=f"remove_{skill}"):
                            skills_to_delete.append(skill)
                    
                    # Experience text area
                    new_experience = st.text_area(
                        f"Your experience with {skill}", 
                        value=data["experience"],
                        key=f"exp_{skill}"
                    )
                    # Update the experience if changed
                    if new_experience != data["experience"]:
                        st.session_state.user_skills[skill]["experience"] = new_experience
                        st.session_state.skills_updated = True
                    
                    st.write("---")
                
                # Process any skills to delete
                for skill in skills_to_delete:
                    del st.session_state.user_skills[skill]
                    st.session_state.skills_updated = True
                    st.success(f"Removed skill: {skill}")
                    st.rerun()
                
                # Add button to save all changes
                if st.session_state.skills_updated:
                    if st.button("Save All Changes"):
                        try:
                            from database import save_user_skills
                            save_result = save_user_skills(st.session_state.user_skills)
                            if save_result:
                                st.session_state.skills_updated = False
                                st.success("All changes saved successfully!")
                            else:
                                st.error("Failed to save changes.")
                        except Exception as e:
                            st.error(f"Error saving changes: {str(e)}")
    
    with graph_tab:
        # No header, just the content
        st.write("Visualize your skills in relation to job requirements from the database.")
        
        if "user_skills" not in st.session_state or not st.session_state.user_skills:
            st.warning("Please add skills in the Skills Profile tab first.")
        else:
            # Set default parameters since we removed the sidebar filters
            skill_type_param = None
            category_param = None
            top_n = 50  # Default number of skills to show
            
            # Get job skills from database
            from database import fetch_job_skills
            
            # Check if we have job postings in session state
            has_job_postings = "job_postings" in st.session_state and len(st.session_state.job_postings) > 0
            
            # Fetch job skills for visualization
            job_skills = fetch_job_skills(top_n=top_n, skill_type=skill_type_param, category=category_param)
            
            # We'll also use skills from session state job postings
            session_job_skills = []
            if has_job_postings:
                # Extract skills from job postings in session state
                for job in st.session_state.job_postings:
                    if "required_skills" in job:
                        for skill in job["required_skills"]:
                            session_job_skills.append({
                                "name": skill,
                                "frequency": 5,
                                "job_count": 1,
                                "skill_type": "required",
                                "job_id": job.get("title", "Unknown Job")
                            })
                    if "preferred_skills" in job:
                        for skill in job["preferred_skills"]:
                            session_job_skills.append({
                                "name": skill,
                                "frequency": 3,
                                "job_count": 1,
                                "skill_type": "preferred",
                                "job_id": job.get("title", "Unknown Job")
                            })
            
            # Combine db job skills with session job skills
            all_job_skills = job_skills + session_job_skills
            
            if not all_job_skills and not has_job_postings:
                st.info("No job skills found. Try uploading some job postings first.")
            else:
                # Create the skill graph - use the combined skills
                graph = generate_skill_graph(st.session_state.user_skills, all_job_skills)
                
                # Create the interactive visualization
                net = create_interactive_graph(graph)
                
                # Get the HTML representation
                html = get_html_network(net)
                
                # Display the visualization
                st.components.v1.html(html, height=600)
                
                # Display skill gap analysis
                missing_skills = identify_skill_gaps(st.session_state.user_skills, job_skills)
                
                st.write("## Skill Gap Analysis")
                st.write("""
                These are the most in-demand skills from job postings that are missing from your profile.
                Focus on acquiring these skills to improve your job market value.
                """)
                
                if missing_skills:
                    import pandas as pd
                    
                    # Create a table of the top missing skills
                    missing_df = pd.DataFrame(
                        [(skill['name'], skill['frequency'], skill.get('skill_type', 'general')) 
                         for skill in missing_skills[:10]], 
                        columns=["Skill", "Job Count", "Type"]
                    )
                    
                    st.table(missing_df)
                    
                    # Button to add all missing skills to profile
                    if st.button("Add Missing Skills to Profile", key="add_graph_skills_btn"):
                        added_count = 0
                        for skill in missing_skills[:10]:
                            skill_name = skill['name']
                            if skill_name not in st.session_state.user_skills:
                                st.session_state.user_skills[skill_name] = {
                                    "rating": 1,
                                    "experience": f"Added from skill gap analysis",
                                    "projects": []
                                }
                                added_count += 1
                        
                        # Save to database
                        try:
                            from database import save_user_skills
                            save_user_skills(st.session_state.user_skills)
                            st.success(f"Added {added_count} new skills to your profile!")
                        except Exception as e:
                            st.warning(f"Could not save skills to database: {e}")
                else:
                    st.success("You have all the skills required by the job postings in the database!")
    
    with roadmap_tab:
        # No header, just the content
        st.write("Plan your learning journey to acquire important skills for your career.")
        
        if "user_skills" not in st.session_state or not st.session_state.user_skills:
            st.warning("Please add skills in the Skills Profile tab first.")
        else:
            # Get job skills from database
            from database import fetch_job_skills
            job_skills = fetch_job_skills(top_n=100)
            
            if not job_skills:
                st.info("No job skills found in the database. Try uploading some job postings first.")
            else:
                # Identify skill gaps
                missing_skills = identify_skill_gaps(st.session_state.user_skills, job_skills)
                
                # Always include all job skills in the options for better selection
                # This ensures all skills from uploaded job postings are available
                skill_options = list(set(s['name'] for s in job_skills))
                
                # Add success message if no missing skills
                if not missing_skills:
                    st.success("You already have all the skills found in job postings! Consider exploring more advanced skills.")
                
                # Let user select multiple skills to learn
                selected_skills = st.multiselect(
                    "Select skills you want to learn",
                    options=skill_options
                )
                
                # Create a list of all skills with their frequency information for the AI
                all_job_skills_data = {}
                for skill in job_skills:
                    name = skill.get('name', '')
                    if name:
                        if name not in all_job_skills_data:
                            all_job_skills_data[name] = {
                                'frequency': 0,
                                'skill_type': skill.get('skill_type', 'general')
                            }
                        all_job_skills_data[name]['frequency'] += 1
                
                # Sort skills by frequency for importance ranking
                ranked_skills = sorted(
                    [(name, data['frequency'], data['skill_type']) 
                     for name, data in all_job_skills_data.items()],
                    key=lambda x: x[1],
                    reverse=True
                )
                
                # Display AI-generated learning path option
                st.subheader("Generate AI Learning Path")
                st.write("Use AI to create a personalized learning path based on in-demand skills and your current skill set.")
                
                # Allow user to select skills or use the multiselect
                if selected_skills:
                    st.success(f"You've selected {len(selected_skills)} skills for your learning path.")
                
                generate_button = st.button("Generate Optimized Learning Path", key="generate_ai_path")
                
                # Add a button to auto-select top skills from job postings
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("Auto-select Top 3 In-demand Skills", key="auto_select_skills"):
                        if missing_skills:
                            # Get top 3 missing skills for auto-selection
                            top_skills = [s['name'] for s in missing_skills[:3]]
                            # Update selected skills in session state to preserve across reruns
                            if 'selected_roadmap_skills' not in st.session_state:
                                st.session_state.selected_roadmap_skills = top_skills
                            else:
                                st.session_state.selected_roadmap_skills = list(set(st.session_state.selected_roadmap_skills + top_skills))
                            st.rerun()  # Rerun to update the multiselect
                
                # Store selected skills in session state
                if selected_skills:
                    st.session_state.selected_roadmap_skills = selected_skills
                elif 'selected_roadmap_skills' in st.session_state:
                    # If selected_skills is empty but we have stored skills, use those
                    selected_skills = st.session_state.selected_roadmap_skills
                
                with col2:
                    if generate_button:
                        with st.spinner("Creating your personalized skill roadmap..."):
                            # Get skills for analysis - either selected or top skills if none selected
                            target_skills = selected_skills
                            if not target_skills and missing_skills:
                                target_skills = [s['name'] for s in missing_skills[:5]]
                                st.info(f"Using top 5 in-demand skills since none were specifically selected.")
                            
                            if not target_skills:
                                st.warning("No skills identified for learning. Please select skills or upload job postings first.")
                            else:
                                try:
                                    # Create a prompt for the OpenAI API with comprehensive context
                                    skills_importance = []
                                    for name, freq, skill_type in ranked_skills[:20]:  # Top 20 skills for context
                                        importance = "required" if skill_type == "required" else "preferred" if skill_type == "preferred" else "general"
                                        skills_importance.append(f"{name} (appears in {freq} job postings, {importance})")
                                    
                                    # Format current skills with ratings
                                    current_skills = []
                                    for skill, data in st.session_state.user_skills.items():
                                        rating = data.get('rating', 3)
                                        current_skills.append(f"{skill} (proficiency: {rating}/5)")
                                    
                                    # Get target skills with importance info
                                    target_skills_with_info = []
                                    for skill in target_skills:
                                        # Find skill in ranked list
                                        skill_found = False
                                        for name, freq, skill_type in ranked_skills:
                                            if name.lower() == skill.lower():
                                                importance = "required" if skill_type == "required" else "preferred" if skill_type == "preferred" else "general"
                                                target_skills_with_info.append(f"{name} (appears in {freq} job postings, {importance})")
                                                skill_found = True
                                                break
                                        if not skill_found:
                                            target_skills_with_info.append(skill)
                                
                                    # Create the prompt for the AI model
                                    prompt = f"""
                                    You are an expert career and skills development advisor creating a personalized learning plan.
                                    
                                    ## Job Market Context
                                    These are the most in-demand skills from job postings, in order of importance:
                                    {', '.join(skills_importance)}
                                    
                                    ## Current Skills
                                    The person has these current skills and proficiency levels:
                                    {', '.join(current_skills)}
                                    
                                    ## Target Skills To Learn
                                    The person wants to learn these skills (with their job market importance):
                                    {', '.join(target_skills_with_info)}
                                    
                                    ## Your Task
                                    Create an intelligent learning roadmap that:
                                    1. Analyzes the relationship between current skills and target skills
                                    2. Identifies any intermediate skills needed for an efficient learning path
                                    3. Prioritizes skills based on job market demand and learning efficiency
                                    4. Suggests a clear sequence that minimizes learning time while maximizing job market value
                                    5. Includes specific learning resources (courses, books, projects)
                                    6. Provides a realistic timeline
                                    
                                    Use the skill graph concept to identify logical skill relationships and progressions.
                                    For example, if they know Python and want to learn machine learning, data analysis
                                    might be an intermediate skill to learn first.
                                    
                                    Format your response with clear headings, bullet points, and a visual learning path
                                    if possible using markdown.
                                    """
                                
                                    # Make the API call
                                    from openai import OpenAI
                                    import os
                                    
                                    client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
                                    
                                    response = client.chat.completions.create(
                                        model="gpt-4o",
                                        messages=[
                                            {"role": "system", "content": "You are a career development and learning specialist who creates detailed skill acquisition roadmaps."},
                                            {"role": "user", "content": prompt}
                                        ]
                                    )
                                    
                                    if response and response.choices:
                                        roadmap_content = response.choices[0].message.content
                                        
                                        # Display the generated roadmap
                                        st.markdown("### Your Personalized Skill Development Plan")
                                        st.markdown(roadmap_content)
                                        
                                        # Add option to save this roadmap using session state to avoid resetting the page
                                        col1, col2 = st.columns([3, 1])
                                        with col2:
                                            save_key = f"save_plan_{datetime.now().strftime('%Y%m%d%H%M%S')}"
                                            save_button = st.button("Save This Learning Plan", key=save_key)
                                        
                                        # Store the roadmap content in session state for preservation through reruns
                                        temp_roadmap_key = "temp_roadmap_content"
                                        if temp_roadmap_key not in st.session_state:
                                            st.session_state[temp_roadmap_key] = roadmap_content
                                        
                                        # Check if the save button was clicked
                                        if save_button:
                                            if "saved_roadmaps" not in st.session_state:
                                                st.session_state.saved_roadmaps = {}
                                            
                                            # Create a name for the roadmap based on selected skills
                                            skills_summary = ", ".join(target_skills[:2])
                                            if len(target_skills) > 2:
                                                skills_summary += f" +{len(target_skills)-2} more"
                                                
                                            # Create a unique timestamp-based name to avoid conflicts
                                            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                                            roadmap_name = f"Skill Plan: {skills_summary} ({timestamp})"
                                                
                                            st.session_state.saved_roadmaps[roadmap_name] = {
                                                "content": roadmap_content,
                                                "date_created": datetime.now().strftime("%Y-%m-%d"),
                                                "skills": target_skills
                                            }
                                            
                                            st.success(f"Learning plan saved successfully!")
                                    else:
                                        st.error("Could not generate learning plan. Please try again later.")
                                except Exception as e:
                                    st.error(f"Error generating learning plan: {str(e)}")
                                    st.info("Please check your internet connection or try again later.")
                
                # Always display the saved plans section
                st.subheader("Your Saved Learning Plans")
                
                if "saved_roadmaps" in st.session_state and st.session_state.saved_roadmaps:
                    # Display all saved plans in reverse chronological order (newest first)
                    plans = list(st.session_state.saved_roadmaps.items())
                    # Sort by timestamp in the name (newest first)
                    plans.sort(key=lambda x: x[0], reverse=True)
                    
                    for plan_name, plan_data in plans:
                        with st.expander(f"{plan_name} (Created: {plan_data['date_created']})"):
                            st.markdown(plan_data["content"])
                            
                            # Show delete button with unique key
                            delete_key = f"del_plan_{plan_name}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
                            if st.button("Delete This Plan", key=delete_key):
                                # Use session state to mark this plan for deletion
                                if "plans_to_delete" not in st.session_state:
                                    st.session_state.plans_to_delete = []
                                st.session_state.plans_to_delete.append(plan_name)
                                st.rerun()
                    
                    # Process any plans marked for deletion
                    if "plans_to_delete" in st.session_state and st.session_state.plans_to_delete:
                        for plan_name in st.session_state.plans_to_delete:
                            if plan_name in st.session_state.saved_roadmaps:
                                del st.session_state.saved_roadmaps[plan_name]
                        st.session_state.plans_to_delete = []
                        st.success("Learning plan deleted successfully!")
                else:
                    st.info("You haven't saved any learning plans yet. Generate a plan and save it to see it here.")
    
    # End of render_skill_graph_tab function


def skill_graph_page():
    """Streamlit page for the interactive skill graph and roadmap visualizations."""
    st.title("Career Skill Graph & Roadmap")
    
    # Initialize session state for user data
    if "user_skills" not in st.session_state:
        # Try to load user skills from database first
        db_skills = fetch_user_skills()
        if db_skills:
            st.session_state.user_skills = db_skills
            print(f"Loaded {len(db_skills)} skills from database")
        else:
            st.session_state.user_skills = {}
    
    if "user_projects" not in st.session_state:
        st.session_state.user_projects = []
        
    # Create tab for rendering the skill graph
    render_skill_graph_tab()
    
    # Save any skills updates to the database if needed
    if st.session_state.get('skills_updated', False):
        try:
            success = save_user_skills(st.session_state.user_skills)
            if success:
                st.session_state.skills_updated = False
                # Only show success message if this isn't the first load
                if not st.session_state.get('is_first_load', True):
                    st.success("Skills saved to database successfully!")
                st.session_state.is_first_load = False
            else:
                st.warning("Failed to save skills to database.")
        except Exception as e:
            st.error(f"Error saving skills to database: {e}")


if __name__ == "__main__":
    skill_graph_page()
